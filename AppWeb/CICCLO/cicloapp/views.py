from django.shortcuts import render, redirect, HttpResponseRedirect
from django.db.models import Count
from django.conf import settings
from django.apps import apps
from django.http import JsonResponse
from django.http import HttpResponse
from .forms import ExcelUploadForm
from .models import *
import matplotlib
matplotlib.use('Agg')
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
import io
from io import BytesIO
import os
import urllib, base64, urllib.parse
import pandas as pd
import seaborn as sns
import numpy as np
from kmodes.kmodes import KModes
import nltk
nltk.download('punkt')
nltk.download('stopwords')
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer
import torch
from transformers import BertTokenizer, BertForSequenceClassification, BertModel
from collections import Counter
from nltk.util import ngrams
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from transformers import BertModel, BertTokenizer
import torch
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import KMeans
import google.generativeai as genai
from IPython.display import Markdown
import textwrap 
from textwrap import wrap
import sys
import re
from docx import Document
from docx.shared import Inches













def home(request):
    return render(request, 'base_generic.html')

def handle_uploaded_file(f):
    # Obtener la ruta del directorio media
    media_root = os.path.join(settings.BASE_DIR, 'media')
    # Crear el directorio uploads si no existe
    if not os.path.exists(os.path.join(media_root, 'uploads')):
        os.makedirs(os.path.join(media_root, 'uploads'))
    # Guardar el archivo en el sistema de archivos del servidor
    file_path = os.path.join(media_root, 'uploads', f.name)
    with open(file_path, 'wb+') as destination:
        for chunk in f.chunks():
            destination.write(chunk)
    return file_path

column_mappingk = {
    "Marca temporal": "marca_temporal",
    "Nombre de usuario": "nombre_usuario",
    "Edad": "edad",
    "Género": "genero",
    "Mencione en qué área de la empresa se desempeña actualmente": "area_empresa",
    "Indique su antigüedad dentro de la empresa": "antiguedad_empresa",
    "1. La normas que rigen la empresa admiten la expresión de la forma de ser de sus empleados.": "pregunta_1",
    "2. Los empleados contribuyen con ideas en la toma de decisiones de la empresa.": "pregunta_2",
    "3. A usted le interesa participar en la toma de las decisiones de la empresa.": "pregunta_3",
    "4. En la empresa cuando se crea una norma, previamente las directivas hacen consultas con los empleados.": "pregunta_4",
    "5. En la empresa a algunas personas les aplican las normas con bastante rigor mientras a otras les perdonan todo.": "pregunta_5",
    "6. En general, la empresa está mejorando en relación a como era cuando usted ingresó como empleado.": "pregunta_6",
    "7. En relación con el día de su ingreso como empleado, usted nota mejoría en el desempeño de los empleados de la empresa.": "pregunta_7",
    "8. La comunicación de trabajo, desde su jefe inmediato hacia usted es fácil.": "pregunta_8",
    "9. La comunicación de trabajo, desde usted hacia su jefe inmediato es fácil.": "pregunta_9",
    "10. La comunicación con los grupos de trabajo con los que usted necesita relacionarse es fácil.": "pregunta_10",
    "11. Como impresión general, usted considera que en la empresa los empleados conocen sus funciones.": "pregunta_11",
    "12. Normalmente la cantidad de trabajo que tiene su cargo es excesiva.": "pregunta_12",
    "13. Las metas que se proponen en la empresa se cumplen.": "pregunta_13",
    "14. Actualmente hay la tendencia en la empresa a desperdiciar insumos de trabajo.": "pregunta_14",
    "15. En la empresa los problemas entre las personas se resuelven fácilmente.": "pregunta_15",
    "16. La forma como está organizado la empresa, es fácil de entender.": "pregunta_16",
    "17. Las tareas son supervisadas excesivamente.": "pregunta_17",
    "18. En la empresa las relaciones entre las personas son cordiales.": "pregunta_18",
    "19. Al interior de la empresa permanentemente hay conflictos.": "pregunta_19",
    "20. Los empleados son solidarios entre sí.": "pregunta_20",
    "21. Las personas en la empresa son tolerantes.": "pregunta_21",
    "22. Los empleados en la empresa se actualizan en los temas que necesita la organización.": "pregunta_22",
    "23. La empresa apoya la autonomía de sus empleados.": "pregunta_23",
    "24. La empresa apoya el desarrollo de carrera (ascensos) de sus empleados.": "pregunta_24",
    "25. La empresa apoya las sugerencias de los empleados para innovar (en procesos, productos, servicios, etc.).": "pregunta_25",
    "26. En la empresa la libertad de expresión se respeta.": "pregunta_26",
    "27. En general, usted se siente bien trabajando en la dependencia actual.": "pregunta_27",
    "28. Usted se siente bien trabajando en la empresa, en general.": "pregunta_28",
    "29. En general, la empresa paga los salarios que cada quien se merece.": "pregunta_29",
    "30. Frente a entidades parecidas, la empresa es fuerte.": "pregunta_30",
    "31. Esta organización le cumple a sus clientes.": "pregunta_31",
    "32. Si usted recibiera una oferta de trabajo de otra organización se iría, siendo las condiciones de la otra las mismas.": "pregunta_32",
    "33. Si usted recibiera una oferta de trabajo de otra organización se iría, siendo las condiciones de la otra mucho mejores.": "pregunta_33",
    "34. Las condiciones de su sitio de trabajo son adecuadas para desempeñarse bien.": "pregunta_34",
    "35. Es notable la presencia de grupos cerrados en los cuales se refugian sus integrantes.": "pregunta_35",
    "36. La cantidad de tareas que tiene su cargo es mayor a la de otros cargos que se le parecen.": "pregunta_36",
    "37. Su jefe sabe cómo hacer el trabajo de sus subalternos.": "pregunta_37",
    "38. Su jefe sabe cómo premiar a sus subalternos.": "pregunta_38",
    "39. Su jefe sabe cómo sancionar a sus subalternos.": "pregunta_39",
    "40. Su puesto de trabajo tiene variedad en la forma de ejecutar las tareas.": "pregunta_40",
    "41. Usted encuentra congruencia entre lo que busca en su vida laboral y lo que le ofrece su puesto de trabajo.": "pregunta_41",
    "42. Cite dos (2) situaciones, anécdotas, historias internas o algo típico, que refleje lo que distingue la cultura de esta organización frente a las que se le parecen. Algo que permita decir: ""esto solo pasa aquí"".\nSituación 1:": "pregunta_42_situacion_1",
    "42. Cite dos (2) situaciones, anécdotas, historias internas o algo típico, que refleje lo que distingue la cultura de esta organización frente a las que se le parecen. Algo que permita decir: ""esto solo pasa aquí"".\nSituación 2:": "pregunta_42_situacion_2",
    "43. Ordene, de mayor a menor importancia, las tres siguientes razones por las cuales usted trabaja aquí (seleccione 3 en la más importante, 2 en la que sigue y 1 en la menos importante): [Me siento bien con mis compañeros]": "pregunta_43_opcion_1",
    "43. Ordene, de mayor a menor importancia, las tres siguientes razones por las cuales usted trabaja aquí (seleccione 3 en la más importante, 2 en la que sigue y 1 en la menos importante): [Puedo ayudar a organizar los equipos de trabajo]": "pregunta_43_opcion_2",
    "43. Ordene, de mayor a menor importancia, las tres siguientes razones por las cuales usted trabaja aquí (seleccione 3 en la más importante, 2 en la que sigue y 1 en la menos importante): [Puedo avanzar hacia las metas que me he propuesto en la vida]": "pregunta_43_opcion_3",
    "44. Indique cuál de las siguientes fuentes de poder tiene mayor influencia en esta entidad (seleccione 3 en la más importante y 1 en la menos importante, seleccione 0 si no aplica): [Las directivas]": "pregunta_44_opcion_1",
    "44. Indique cuál de las siguientes fuentes de poder tiene mayor influencia en esta entidad (seleccione 3 en la más importante y 1 en la menos importante, seleccione 0 si no aplica): [Los empleados]": "pregunta_44_opcion_2",
    "44. Indique cuál de las siguientes fuentes de poder tiene mayor influencia en esta entidad (seleccione 3 en la más importante y 1 en la menos importante, seleccione 0 si no aplica): [Factores externos a la entidad]": "pregunta_44_opcion_3",
    "44. Indique cuál de las siguientes fuentes de poder tiene mayor influencia en esta entidad (seleccione 3 en la más importante y 1 en la menos importante, seleccione 0 si no aplica): [Factores externos a la entidad]": "pregunta_44_opcion_4",
    "45. Mencione tres (3) defectos de esta entidad.\nDefecto 1:": "pregunta_45_defecto_1",
    "45. Mencione tres (3) defectos de esta entidad.\nDefecto 2:": "pregunta_45_defecto_2",
    "45. Mencione tres (3) defectos de esta entidad.\nDefecto 3:": "pregunta_45_defecto_3",
    "46. Mencione tres (3) virtudes de esta entidad.\nVirtud 1:": "pregunta_46_virtud_1",
    "46. Mencione tres (3) virtudes de esta entidad.\nVirtud 2:": "pregunta_46_virtud_2",
    "46. Mencione tres (3) virtudes de esta entidad.\nVirtud 3:": "pregunta_46_virtud_3",
    "47. Mencione dos (2) hábitos diarios que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 1:": "pregunta_47_habito_diario_1",
    "47. Mencione dos (2) hábitos diarios que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 2:": "pregunta_47_habito_diario_2",
    "48. Mencione dos (2) hábitos mensuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 1:": "pregunta_48_habito_mensual_1",
    "48. Mencione dos (2) hábitos mensuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 2:": "pregunta_48_habito_mensual_2",
    "49. Mencione dos (2) hábitos anuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 1:": "pregunta_49_habito_anual_1",
    "49. Mencione dos (2) hábitos anuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados.\nHábito 2:": "pregunta_49_habito_anual_2",
    "50. Piense en tres personas que se destacan negativamente dentro de la empresa y señale sólo sus defectos. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nDefectos persona A:": "pregunta_50_defecto_persona_A",
    "50. Piense en tres personas que se destacan negativamente dentro de la empresa y señale sólo sus defectos. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nDefectos persona B:": "pregunta_50_defecto_persona_B",
    "50. Piense en tres personas que se destacan negativamente dentro de la empresa y señale sólo sus defectos. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nDefectos persona C:": "pregunta_50_defecto_persona_C",
    "51. Piense en tres personas que se destacan positivamente dentro de la empresa y señale sólo sus virtudes. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nVirtudes persona A:": "pregunta_51_virtud_persona_A",
    "51. Piense en tres personas que se destacan positivamente dentro de la empresa y señale sólo sus virtudes. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nVirtudes persona B:": "pregunta_51_virtud_persona_B",
    "51. Piense en tres personas que se destacan positivamente dentro de la empresa y señale sólo sus virtudes. No diga los nombres, sólo recuérdelas como las personas A, B y C.\nVirtudes persona C:": "pregunta_51_virtud_persona_C",
    "52. A cuál de sus compañeros elegiría para que organice una fiesta de integración en la empresa.": "pregunta_52",
    "53. A quién de la empresa elegiría para que defienda los intereses de su grupo profesional.": "pregunta_53",
    "54. A quién de la empresa elegiría para que lo represente ante las directivas de esta organización.": "pregunta_54",
    "55. A quién dentro de la empresa elegiría para que organice un equipo deportivo.": "pregunta_55",
    "56. A quién dentro de la empresa elegiría para que organice los equipos de trabajo.": "pregunta_56",
    "57. A quién dentro de la empresa elegiría para comentar y divulgar los hechos de la vida cotidiana de la organización.": "pregunta_57",
    "58. A cuál de sus compañeros le confiaría un secreto.": "pregunta_58",
    "59. A quién dentro de la empresa elegiría para resolver problemas entre compañeros de trabajo.": "pregunta_59",
    "60. A cuál de sus compañeros elegiría para que le enseñara a mejorar la forma de hacer su trabajo.": "pregunta_60",
    "61. Mencione a un funcionario de la empresa que según usted tiene rasgos de líder.": "pregunta_61",
    "62. Coloque 3 en lo que su jefe hace con más frecuencia y 1 en lo que casi nunca hace: (Primero leerle las 3 opciones completas. Y después leérselas una por una, para que las ordene). [Que los subalternos se sientan bien en sus sitios de trabajo aunque no": "pregunta_62_opcion_1",
    "62. Coloque 3 en lo que su jefe hace con más frecuencia y 1 en lo que casi nunca hace: (Primero leerle las 3 opciones completas. Y después leérselas una por una, para que las ordene). [Que las tareas se hagan bien y que los empleados estén bien]": "pregunta_62_opcion_2",
    "62. Coloque 3 en lo que su jefe hace con más frecuencia y 1 en lo que casi nunca hace: (Primero leerle las 3 opciones completas. Y después leérselas una por una, para que las ordene). [Que las tareas se hagan bien aunque los empleados estén mal]": "pregunta_62_opcion_3",
    "Puede escribir algunos comentarios adicionales si lo desea.": "comentarios",
}

column_mapping = [
    "marca_temporal",
    "nombre_usuario",
    "edad",
    "genero",
    "area_empresa",
    "antiguedad_empresa",
    "pregunta_1",
    "pregunta_2",
    "pregunta_3",
    "pregunta_4",
    "pregunta_5",
    "pregunta_6",
    "pregunta_7",
    "pregunta_8",
    "pregunta_9",
    "pregunta_10",
    "pregunta_11",
    "pregunta_12",
    "pregunta_13",
    "pregunta_14",
    "pregunta_15",
    "pregunta_16",
    "pregunta_17",
    "pregunta_18",
    "pregunta_19",
    "pregunta_20",
    "pregunta_21",
    "pregunta_22",
    "pregunta_23",
    "pregunta_24",
    "pregunta_25",
    "pregunta_26",
    "pregunta_27",
    "pregunta_28",
    "pregunta_29",
    "pregunta_30",
    "pregunta_31",
    "pregunta_32",
    "pregunta_33",
    "pregunta_34",
    "pregunta_35",
    "pregunta_36",
    "pregunta_37",
    "pregunta_38",
    "pregunta_39",
    "pregunta_40",
    "pregunta_41",
    "pregunta_42_situacion_1",
    "pregunta_42_situacion_2",
    "pregunta_43_opcion_1",
    "pregunta_43_opcion_2",
    "pregunta_43_opcion_3",
    "pregunta_44_opcion_1",
    "pregunta_44_opcion_2",
    "pregunta_44_opcion_3",
    "pregunta_44_opcion_4",
    "pregunta_45_defecto_1",
    "pregunta_45_defecto_2",
    "pregunta_45_defecto_3",
    "pregunta_46_virtud_1",
    "pregunta_46_virtud_2",
    "pregunta_46_virtud_3",
    "pregunta_47_habito_diario_1",
    "pregunta_47_habito_diario_2",
    "pregunta_48_habito_mensual_1",
    "pregunta_48_habito_mensual_2",
    "pregunta_49_habito_anual_1",
    "pregunta_49_habito_anual_2",
    "pregunta_50_defecto_persona_A",
    "pregunta_50_defecto_persona_B",
    "pregunta_50_defecto_persona_C",
    "pregunta_51_virtud_persona_A",
    "pregunta_51_virtud_persona_B",
    "pregunta_51_virtud_persona_C",
    "pregunta_52",
    "pregunta_53",
    "pregunta_54",
    "pregunta_55",
    "pregunta_56",
    "pregunta_57",
    "pregunta_58",
    "pregunta_59",
    "pregunta_60",
    "pregunta_61",
    "pregunta_62_opcion_1",
    "pregunta_62_opcion_2",
    "pregunta_62_opcion_3",
    "comentarios",
]


def upload_file(request):               
    if request.method == 'POST':
        form = ExcelUploadForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            df = pd.read_excel(excel_file)
            # Renombrar las columnas del DataFrame utilizando column_mapping
            df.columns = column_mapping
            # Convertir la columna 'marca_temporal' a un formato válido
            df['marca_temporal'] = pd.to_datetime(df['marca_temporal'])



            # Iterar sobre el DataFrame y guardar los datos en el modelo DatosDemograficos
            for index, row in df.iterrows():
                nuevo_registro = DatosDemograficos(
                    marca_temporal=row['marca_temporal'],
                    nombre_usuario=row['nombre_usuario'],
                    edad=row['edad'],
                    genero=row['genero'],
                    area_empresa=row['area_empresa'],
                    antiguedad_empresa=row['antiguedad_empresa']
                )
                nuevo_registro.save()



            # Iterar solo sobre las columnas específicas que representan preguntas cerradas de opción múltiple (preguntas 1 a 41)
            for index, row in df.iterrows():
                nueva_pregunta = PreguntasCerradas(
                        pregunta_1=row['pregunta_1'],
                        pregunta_2=row['pregunta_2'],
                        pregunta_3=row['pregunta_3'],
                        pregunta_4=row['pregunta_4'],
                        pregunta_5=row['pregunta_5'],
                        pregunta_6=row['pregunta_6'],
                        pregunta_7=row['pregunta_7'],
                        pregunta_8=row['pregunta_8'],
                        pregunta_9=row['pregunta_9'],
                        pregunta_10=row['pregunta_10'],
                        pregunta_11=row['pregunta_11'],
                        pregunta_12=row['pregunta_12'],
                        pregunta_13=row['pregunta_13'],
                        pregunta_14=row['pregunta_14'],
                        pregunta_15=row['pregunta_15'],
                        pregunta_16=row['pregunta_16'],
                        pregunta_17=row['pregunta_17'],
                        pregunta_18=row['pregunta_18'],
                        pregunta_19=row['pregunta_19'],
                        pregunta_20=row['pregunta_20'],
                        pregunta_21=row['pregunta_21'],
                        pregunta_22=row['pregunta_22'],
                        pregunta_23=row['pregunta_23'],
                        pregunta_24=row['pregunta_24'],
                        pregunta_25=row['pregunta_25'],
                        pregunta_26=row['pregunta_26'],
                        pregunta_27=row['pregunta_27'],
                        pregunta_28=row['pregunta_28'],
                        pregunta_29=row['pregunta_29'],
                        pregunta_30=row['pregunta_30'],
                        pregunta_31=row['pregunta_31'],
                        pregunta_32=row['pregunta_32'],
                        pregunta_33=row['pregunta_33'],
                        pregunta_34=row['pregunta_34'],
                        pregunta_35=row['pregunta_35'],
                        pregunta_36=row['pregunta_36'],
                        pregunta_37=row['pregunta_37'],
                        pregunta_38=row['pregunta_38'],
                        pregunta_39=row['pregunta_39'],
                        pregunta_40=row['pregunta_40'],
                        pregunta_41=row['pregunta_41']
                    )
                nueva_pregunta.save()



            # Iterar sobre las columnas especificadas que representan la pregunta 42
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbierta
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbierta.objects.create(
                    pregunta_42_situacion_1=row['pregunta_42_situacion_1'],
                    pregunta_42_situacion_2=row['pregunta_42_situacion_2']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas especificadas que representan las preguntas 43 y 44
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaImportancia
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaImportancia.objects.create(
                    pregunta_43_opcion_1=row['pregunta_43_opcion_1'],
                    pregunta_43_opcion_2=row['pregunta_43_opcion_2'],
                    pregunta_43_opcion_3=row['pregunta_43_opcion_3'],
                    pregunta_44_opcion_1=row['pregunta_44_opcion_1'],
                    pregunta_44_opcion_2=row['pregunta_44_opcion_2'],
                    pregunta_44_opcion_3=row['pregunta_44_opcion_3'],
                    pregunta_44_opcion_4=row['pregunta_44_opcion_4']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan los defectos de la pregunta 45
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaDefectos
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaDefectos.objects.create(
                    defecto_1=row['pregunta_45_defecto_1'],
                    defecto_2=row['pregunta_45_defecto_2'],
                    defecto_3=row['pregunta_45_defecto_3']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan las virtudes de la pregunta 46
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaVirtudes
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaVirtudes.objects.create(
                    virtud_1=row['pregunta_46_virtud_1'],
                    virtud_2=row['pregunta_46_virtud_2'],
                    virtud_3=row['pregunta_46_virtud_3']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan los hábitos diarios de la pregunta 47
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaHabitos
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaHabitos.objects.create(
                    habito_1=row['pregunta_47_habito_diario_1'],
                    habito_2=row['pregunta_47_habito_diario_2']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan los hábitos mensuales de la pregunta 48
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaHabitosMensuales
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaHabitosMensuales.objects.create(
                    habito_1=row['pregunta_48_habito_mensual_1'],
                    habito_2=row['pregunta_48_habito_mensual_2']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan los hábitos anuales de la pregunta 49
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaHabitosAnuales
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaHabitosAnuales.objects.create(
                    habito_1=row['pregunta_49_habito_anual_1'],
                    habito_2=row['pregunta_49_habito_anual_2']
                )
                nueva_pregunta.save()




            # Iterar sobre las columnas específicas que representan los defectos de personas de la pregunta 50
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaDefectosPersonas
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaDefectosPersonas.objects.create(
                    defectos_persona_A=row['pregunta_50_defecto_persona_A'],
                    defectos_persona_B=row['pregunta_50_defecto_persona_B'],
                    defectos_persona_C=row['pregunta_50_defecto_persona_C']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan las virtudes de personas de la pregunta 51
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaVirtudesPersonas
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaVirtudesPersonas.objects.create(
                    virtudes_persona_A=row['pregunta_51_virtud_persona_A'],
                    virtudes_persona_B=row['pregunta_51_virtud_persona_B'],
                    virtudes_persona_C=row['pregunta_51_virtud_persona_C']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan las preguntas abiertas
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaAbiertaCompaneros
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaAbiertaCompaneros.objects.create(
                    fiesta_integracion=row['pregunta_52'],
                    defensa_intereses=row['pregunta_53'],
                    representante_directivas=row['pregunta_54'],
                    organizador_equipo_deportivo=row['pregunta_55'],
                    organizador_equipos_trabajo=row['pregunta_56'],
                    divulgacion_hechos=row['pregunta_57'],
                    confianza_secreto=row['pregunta_58'],
                    resolver_problemas=row['pregunta_59'],
                    enseñanza_trabajo=row['pregunta_60'],
                    lider_funcionario=row['pregunta_61']
                )
                nueva_pregunta.save()



            # Iterar sobre las columnas específicas que representan las opciones de la pregunta 62
            # Iterar sobre el DataFrame y guardar los datos en el modelo PreguntaOrden
            for index, row in df.iterrows():
                nueva_pregunta = PreguntaOrden.objects.create(
                    pregunta_62_opcion_1=row['pregunta_62_opcion_1'],
                    pregunta_62_opcion_2=row['pregunta_62_opcion_2'],
                    pregunta_62_opcion_3=row['pregunta_62_opcion_3']
                )
                nueva_pregunta.save()


            #Comentarios Adicionales
            for index, row in df.iterrows():
                comentarios = row['comentarios']
                if comentarios.strip():  # Verificar si comentarios no está vacío
                    comentarios_adicionales = ComentariosAdicionales.objects.create(
                        comentarios=comentarios
                    )
                    comentarios_adicionales.save()



                # Redirigir al usuario a la página de resultados
                return redirect('view_results')
        else:
            return render(request, 'upload.html', {'form': form})
    else:
        form = ExcelUploadForm()
        return render(request, 'upload.html', {'form': form})





def base_generic(request):
    return render(request, 'base_generic.html')




#Boton para vaciar datos
def delete_data(request):
    # Eliminar todos los objetos de todos los modelos
    for model in apps.get_app_config('cicloapp').get_models():
        model.objects.all().delete()

    # Eliminar todas las gráficas generadas
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    for chart_file in os.listdir(chart_dir):
        chart_path = os.path.join(chart_dir, chart_file)
        if os.path.isfile(chart_path):
            os.remove(chart_path)

    # Redirigir a la vista base_generic
    return redirect('base_generic')








#Boton para crear Word
def export_word(request):
    # Directorio donde se encuentran las gráficas
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    
    # Crear un documento Word nuevo
    document = Document()
    document.add_heading('Gráficas Generadas', 0)
    
    # Recorrer los archivos del directorio de gráficos
    for filename in os.listdir(chart_dir):
        file_path = os.path.join(chart_dir, filename)
        if os.path.isfile(file_path):
            # Agregar un subtítulo con el nombre del archivo
            document.add_heading(filename, level=1)
            try:
                # Agregar la imagen al documento (puedes ajustar el ancho)
                document.add_picture(file_path, width=Inches(6))
            except Exception as e:
                # En caso de error (por ejemplo, si el archivo no es una imagen válida), se omite
                document.add_paragraph(f'No se pudo agregar la imagen: {e}')
    
    # Guardar el documento en un objeto BytesIO
    f = BytesIO()
    document.save(f)
    f.seek(0)
    
    # Configurar la respuesta HTTP para descargar el archivo Word
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="graficas.docx"'
    return response









def view_data(request):
    # Obtener todos los modelos en la aplicación 'cicloapp'
    models = apps.get_app_config('cicloapp').get_models()

    # Crear un diccionario para almacenar los datos de cada modelo
    data = {}

    # Iterar sobre cada modelo
    for model in models:
        # Obtener todos los objetos del modelo actual
        objects = model.objects.all()

        # Obtener los datos como un diccionario
        data_dict = objects.values()

        # Crear un DataFrame especificando los nombres de las columnas
        df = pd.DataFrame(data_dict)

        # Convertir el DataFrame en una lista de diccionarios para poder pasarlo a la plantilla
        data[model.__name__] = df.to_dict(orient='records')


    # Renderizar la plantilla con los datos
    return render(request, 'data.html', {'data': data})








































# Función para generar la gráfica de barras
def generate_age_bar_chart(request):
    # Obtener los datos de la columna edad
    age_data = DatosDemograficos.objects.values('edad').annotate(count=Count('edad')).order_by('edad')

    # Extraer las edades y sus frecuencias
    ages = [item['edad'] for item in age_data]
    frequencies = [item['count'] for item in age_data]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras y la separación entre ellas
    bar_width = 0.8  # Ancho de las barras
    bar_spacing = 0.1  # Espacio entre barras

    # Calcular la ubicación de cada barra
    bar_positions = np.arange(len(ages))

    # Crear la gráfica de barras
    bars = ax.bar(bar_positions, frequencies, width=bar_width, color='skyblue', align='center')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Edad')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Distribución de Edades')
    ax.set_xticks(bar_positions)  # Establecer las posiciones de las barras como marcas en el eje x
    ax.set_xticklabels(ages)  # Establecer las edades como etiquetas en el eje x

    # Mostrar la cantidad exacta de veces que se repite cada edad en el eje y
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{}'.format(height),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica (horizontal)
    description = 'Esta gráfica muestra la distribución de edades con la frecuencia exacta de cada edad'
    ax.text(1.02, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'age_distribution_chart.png')
    fig.tight_layout(pad=4.0)  # Ajuste automático del diseño para que el texto no se corte
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'age_distribution_chart.png')

    return chart_url  # Devolver la URL de la imagen







def generate_gender_bar_chart(request):
    # Obtener los datos de la columna género
    gender_data = DatosDemograficos.objects.values('genero').annotate(count=Count('genero')).order_by('genero')

    # Extraer los géneros y sus frecuencias
    genders = [item['genero'] for item in gender_data]
    frequencies = [item['count'] for item in gender_data]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras y la separación entre ellas
    bar_width = 0.5  # Ancho de las barras

    # Calcular la ubicación de cada barra
    bar_positions = np.arange(len(genders))

    # Crear la gráfica de barras
    bars = ax.bar(bar_positions, frequencies, width=bar_width, color='lightgreen', align='center')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Género')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Distribución de Género')
    ax.set_xticks(bar_positions)  # Establecer las posiciones de las barras como marcas en el eje x
    ax.set_xticklabels(genders)  # Establecer los géneros como etiquetas en el eje x

    # Mostrar la cantidad exacta de veces que se repite cada género en el eje y
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{}'.format(height),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica (horizontal)
    description = 'Esta gráfica muestra la distribución de género con la frecuencia exacta de cada género'
    ax.text(1.02, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'gender_distribution_chart.png')
    fig.tight_layout(pad=4.0)  # Ajuste automático del diseño para que el texto no se corte
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'gender_distribution_chart.png')

    return chart_url  # Devolver la URL de la imagen










def generate_area_empresa_chart(request):
    # Obtener los datos de la columna "area_empresa"
    area_empresa_data = DatosDemograficos.objects.values('area_empresa').annotate(count=Count('area_empresa')).order_by('area_empresa')

    # Normalizar las áreas de empresa y sus frecuencias
    areas_empresa = [item['area_empresa'] for item in area_empresa_data]
    frequencies = [item['count'] for item in area_empresa_data]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(12, 6))

    # Crear la gráfica de barras
    bars = ax.bar(areas_empresa, frequencies, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Área de Empresa')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Distribución de Áreas de Empresa')

    # Rotar los nombres de las barras y ajustar el espaciado
    ax.set_xticks(range(len(areas_empresa)))  # Establecer las posiciones de las barras como marcas en el eje x
    ax.set_xticklabels(areas_empresa, rotation=45, ha='right')  # Rotar los nombres de las barras

    # Mostrar la cantidad exacta de veces que se repite cada área de empresa en el eje y
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{}'.format(height),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica (horizontal)
    description = 'Esta gráfica muestra la distribución de áreas de empresa con la frecuencia exacta de cada área'
    ax.text(1.02, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'area_empresa_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'area_empresa_chart.png')

    return chart_url









def generate_antiguedad_empresa_chart(request):
    # Obtener los datos de la columna "antiguedad_empresa"
    antiguedad_empresa_data = DatosDemograficos.objects.values('antiguedad_empresa').annotate(count=Count('antiguedad_empresa')).order_by('antiguedad_empresa')

    # Extraer las antigüedades de empresa y sus frecuencias
    antiguedades_empresa = [item['antiguedad_empresa'] for item in antiguedad_empresa_data]
    frequencies = [item['count'] for item in antiguedad_empresa_data]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(12, 6))

    # Crear la gráfica de barras
    bars = ax.bar(antiguedades_empresa, frequencies, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Antigüedad en la Empresa')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Distribución de Antigüedad en la Empresa')

    # Mostrar la cantidad exacta de veces que se repite cada antigüedad en la empresa en el eje y
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{}'.format(height),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica (horizontal)
    description = 'Esta gráfica muestra la distribución de antigüedad en la empresa con la frecuencia exacta de cada nivel'
    ax.text(1.02, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'antiguedad_empresa_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'antiguedad_empresa_chart.png')

    return chart_url







#Monitoreo de datos
def obtener_opciones_y_frecuencias():
    # Obtener todas las instancias del modelo PreguntasCerradas
    preguntas_cerradas = PreguntasCerradas.objects.all()

    # Inicializar un diccionario para almacenar las opciones y sus frecuencias para cada pregunta
    opciones_frecuencias_por_pregunta = {}

    # Iterar sobre cada campo de pregunta en el modelo PreguntasCerradas
    for campo in PreguntasCerradas._meta.fields:
        if campo.name.startswith('pregunta_'):
            # Obtener todas las respuestas para la pregunta actual
            respuestas = [getattr(pregunta, campo.name) for pregunta in preguntas_cerradas]

            # Contar las frecuencias de cada respuesta
            frecuencias = Counter(respuestas)

            # Almacenar las opciones y sus frecuencias en el diccionario
            opciones_frecuencias_por_pregunta[campo.name] = frecuencias

    return opciones_frecuencias_por_pregunta


def tabla_datos(request):
    # Llamar a la función para obtener las opciones y frecuencias
    opciones_frecuencias = obtener_opciones_y_frecuencias()

    # Crear un DataFrame de pandas para mostrar los resultados
    df = pd.DataFrame(opciones_frecuencias).T

    # Calcular el total de encuestados por cada ítem
    df['Total entrevistados'] = df.sum(axis=1)

    # Calcular los porcentajes de cada opción de respuesta para cada ítem
    for columna in df.columns[:-1]:
        df[columna + ' (%)'] = (df[columna] / df['Total entrevistados'] * 100).round(2)  
    # Crear una figura y ejes de tabla con matplotlib.figure.Figure
    fig = Figure(figsize=(30, 20), dpi=100)
    ax = fig.subplots()

    # Ocultar los ejes
    ax.axis('tight')
    ax.axis('off')

    # Crear la tabla en los ejes
    tabla = ax.table(cellText=df.values,
                     colLabels=df.columns,
                     loc='center',
                     cellLoc='center',
                     colWidths=[0.11] * len(df.columns),  # Ajuste de ancho de columnas
                     colColours=["#f5f5f5"] * len(df.columns))

    # Ajustar el diseño de la tabla
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)


    # Crear un lienzo de figura para renderizar la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'tabla_datos.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    table_url = os.path.join(settings.MEDIA_URL, 'charts', 'tabla_datos.png')

    return table_url










def generate_pregunta_1_chart(request):
    # Obtener los datos de la pregunta_1
    pregunta_1_data = PreguntasCerradas.objects.values('pregunta_1').annotate(count=Count('pregunta_1')).order_by('pregunta_1')

    # Extraer las opciones de la pregunta_1 y sus frecuencias
    opciones = [item['pregunta_1'] for item in pregunta_1_data]
    frequencies = [item['count'] for item in pregunta_1_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('1. Las normas que rigen la empresa admiten la expresión de la forma de ser de sus empleados.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"La mayoría de los encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "consideran que las normas empresariales favorecen la expresión personal de los empleados. Esto sugiere un ambiente laboral flexible y abierto."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "perciben que las normas empresariales no permiten la expresión personal. Esto podría reflejar una cultura organizacional más rígida."
        )
    else:
        inferencia = (
            f"Las percepciones sobre la flexibilidad de las normas empresariales están divididas, "
            f"con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos. Esto muestra una polarización en las opiniones."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 1.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_1_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_1_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }











def generate_pregunta_2_chart(request):
    # Obtener los datos de la pregunta_2
    pregunta_2_data = PreguntasCerradas.objects.values('pregunta_2').annotate(count=Count('pregunta_2')).order_by('pregunta_2')

    # Extraer las opciones de la pregunta_2 y sus frecuencias
    opciones = [item['pregunta_2'] for item in pregunta_2_data]
    frequencies = [item['count'] for item in pregunta_2_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('2. Los empleados contribuyen con ideas en la toma de decisiones de la empresa.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"La mayoría de los encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "perciben que los empleados contribuyen con ideas en la toma de decisiones. Esto puede reflejar un estilo de liderazgo participativo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "consideran que los empleados no tienen participación en las decisiones. Esto podría indicar una estructura organizacional más centralizada."
        )
    else:
        inferencia = (
            f"Las opiniones están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere diferentes perspectivas dentro de la organización."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 2.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_2_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_2_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }








def generate_pregunta_3_chart(request):
    # Obtener los datos de la pregunta_3
    pregunta_3_data = PreguntasCerradas.objects.values('pregunta_3').annotate(count=Count('pregunta_3')).order_by('pregunta_3')

    # Extraer las opciones de la pregunta_3 y sus frecuencias
    opciones = [item['pregunta_3'] for item in pregunta_3_data]
    frequencies = [item['count'] for item in pregunta_3_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('3. A usted le interesa participar en la toma de las decisiones de la empresa.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"La mayoría de los encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "están interesados en participar en la toma de decisiones de la empresa, lo que sugiere una inclinación hacia la colaboración y la inclusión."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número significativo de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no están interesados en participar, lo que podría reflejar falta de confianza o preferencia por un enfoque más directivo."
        )
    else:
        inferencia = (
            f"Las opiniones están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que indica una diversidad de perspectivas sobre el interés en participar en las decisiones."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 3.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_3_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_3_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }








def generate_pregunta_4_chart(request):
    # Obtener los datos de la pregunta_4
    pregunta_4_data = PreguntasCerradas.objects.values('pregunta_4').annotate(count=Count('pregunta_4')).order_by('pregunta_4')

    # Extraer las opciones de la pregunta_4 y sus frecuencias
    opciones = [item['pregunta_4'] for item in pregunta_4_data]
    frequencies = [item['count'] for item in pregunta_4_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('4. En la empresa cuando se crea una norma, previamente las directivas hacen consultas con los empleados.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"La mayoría de los encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "perciben que las directivas consultan a los empleados antes de crear normas, reflejando un enfoque participativo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número considerable de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "considera que las directivas no consultan a los empleados, lo que podría señalar falta de comunicación o inclusión."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "sugiriendo una diversidad de opiniones sobre la consulta previa a los empleados."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 4.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_4_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_4_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }









def generate_pregunta_5_chart(request):
    # Obtener los datos de la pregunta_5
    pregunta_5_data = PreguntasCerradas.objects.values('pregunta_5').annotate(count=Count('pregunta_5')).order_by('pregunta_5')

    # Extraer las opciones de la pregunta_5 y sus frecuencias
    opciones = [item['pregunta_5'] for item in pregunta_5_data]
    frequencies = [item['count'] for item in pregunta_5_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('5. En la empresa a algunas personas les aplican las normas con bastante rigor mientras a otras les perdonan todo.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que las normas no se aplican de manera uniforme, lo que podría indicar favoritismos o inconsistencia en la gestión."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "percibe que las normas se aplican de manera justa y uniforme en la organización."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la aplicación de normas en la empresa."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 5.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_5_chart.png')
    try:
        fig.savefig(chart_path)
    except OSError as e:
        print(f"Error saving chart: {e}")
        print(f"Chart Path: {chart_path}")
        # Handle the error, maybe log it, or return an error response
    plt.close(fig)  # Cierra la figura y libera memoria


    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_5_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }











def generate_pregunta_6_chart(request):
    # Obtener los datos de la pregunta_6
    pregunta_6_data = PreguntasCerradas.objects.values('pregunta_6').annotate(count=Count('pregunta_6')).order_by('pregunta_6')

    # Extraer las opciones de la pregunta_6 y sus frecuencias
    opciones = [item['pregunta_6'] for item in pregunta_6_data]
    frequencies = [item['count'] for item in pregunta_6_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('6. En general, la empresa está mejorando en relación a como era cuando usted ingresó como empleado.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que la empresa ha mejorado desde su ingreso, lo que refleja percepciones positivas del progreso organizacional."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe mejoras en la empresa, lo que podría indicar retos en la gestión o comunicación del progreso."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la mejora de la empresa desde el ingreso de los empleados."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 6.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_6_chart.png')
    fig.savefig(chart_path)

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_6_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }









def generate_pregunta_7_chart(request):
    # Obtener los datos de la pregunta_7
    pregunta_7_data = PreguntasCerradas.objects.values('pregunta_7').annotate(count=Count('pregunta_7')).order_by('pregunta_7')

    # Extraer las opciones de la pregunta_7 y sus frecuencias
    opciones = [item['pregunta_7'] for item in pregunta_7_data]
    frequencies = [item['count'] for item in pregunta_7_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('7. En relación con el día de su ingreso como empleado, usted nota mejoría en el desempeño de los empleados de la empresa.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que el desempeño de los empleados ha mejorado, lo que refleja avances positivos en la organización."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe mejoras en el desempeño de los empleados, lo que podría indicar áreas de oportunidad en el desarrollo organizacional."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre el desempeño de los empleados desde el ingreso de los encuestados."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 7.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_7_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_7_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }







def generate_pregunta_8_chart(request):

    # Obtener los datos de la pregunta_8
    pregunta_8_data = PreguntasCerradas.objects.values('pregunta_8').annotate(count=Count('pregunta_8')).order_by('pregunta_8')

    # Extraer las opciones de la pregunta_8 y sus frecuencias
    opciones = [item['pregunta_8'] for item in pregunta_8_data]
    frequencies = [item['count'] for item in pregunta_8_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con todas las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('8. La comunicación de trabajo, desde su jefe inmediato hacia usted es fácil.')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"La mayoría de los encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) perciben que la comunicación "
            f"de trabajo desde su jefe inmediato hacia ellos es positiva. Esto indica una buena percepción de liderazgo y comunicación."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) percibe "
            f"dificultades en la comunicación con su jefe inmediato. Esto podría sugerir áreas de mejora en habilidades de comunicación y liderazgo."
        )
    else:
        inferencia = (
            f"Las percepciones sobre la comunicación con el jefe inmediato están divididas equitativamente, "
            f"con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos. Esto indica una polarización en las experiencias."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 8.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_8_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_8_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }









def generate_pregunta_9_chart(request):
    # Obtener los datos de la pregunta_9
    pregunta_9_data = PreguntasCerradas.objects.values('pregunta_9').annotate(count=Count('pregunta_9')).order_by('pregunta_9')

    # Extraer las opciones de la pregunta_9 y sus frecuencias
    opciones = [item['pregunta_9'] for item in pregunta_9_data]
    frequencies = [item['count'] for item in pregunta_9_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('9. La comunicación de trabajo, desde usted hacia su jefe inmediato es fácil.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que la comunicación hacia su jefe inmediato es fácil, reflejando un entorno positivo para el diálogo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe facilidad en la comunicación con su jefe inmediato, lo que podría señalar desafíos en la interacción jerárquica."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la facilidad de comunicación hacia el jefe inmediato."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 9.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_9_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_9_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }











def generate_pregunta_10_chart(request):
    # Obtener los datos de la pregunta_10
    pregunta_10_data = PreguntasCerradas.objects.values('pregunta_10').annotate(count=Count('pregunta_10')).order_by('pregunta_10')

    # Extraer las opciones de la pregunta_10 y sus frecuencias
    opciones = [item['pregunta_10'] for item in pregunta_10_data]
    frequencies = [item['count'] for item in pregunta_10_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('10. La comunicación con los grupos de trabajo con los que usted necesita relacionarse es fácil.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que la comunicación con los grupos de trabajo es fácil, lo que indica un entorno colaborativo positivo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe facilidad en la comunicación con los grupos de trabajo, lo que podría reflejar desafíos en la colaboración."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la facilidad de comunicación con los grupos de trabajo necesarios."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 10.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_10_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_10_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }











def generate_pregunta_11_chart(request):
    # Obtener los datos de la pregunta_11
    pregunta_11_data = PreguntasCerradas.objects.values('pregunta_11').annotate(count=Count('pregunta_11')).order_by('pregunta_11')

    # Extraer las opciones de la pregunta_11 y sus frecuencias
    opciones = [item['pregunta_11'] for item in pregunta_11_data]
    frequencies = [item['count'] for item in pregunta_11_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('11. Como impresión general, usted considera que en la empresa los empleados conocen sus funciones.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que los empleados conocen sus funciones, lo que indica una claridad organizacional positiva."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe que los empleados conozcan sus funciones, lo que podría reflejar una falta de claridad en los roles laborales."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre el conocimiento de las funciones por parte de los empleados."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 11.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_11_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_11_chart.png')
    
    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }








def generate_pregunta_12_chart(request):
    # Obtener los datos de la pregunta_12
    pregunta_12_data = PreguntasCerradas.objects.values('pregunta_12').annotate(count=Count('pregunta_12')).order_by('pregunta_12')

    # Extraer las opciones de la pregunta_12 y sus frecuencias
    opciones = [item['pregunta_12'] for item in pregunta_12_data]
    frequencies = [item['count'] for item in pregunta_12_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria de acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('12. Los empleados tienen claro a quién reportar en caso de problemas.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que los empleados tienen claro a quién reportar en caso de problemas, lo que indica una buena estructura jerárquica."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe claridad en a quién reportar problemas, lo que podría reflejar una falta de estructura organizativa."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la claridad en la cadena de reporte de problemas."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 12.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_12_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_12_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }














def generate_pregunta_13_chart(request):
    # Obtener los datos de la pregunta_13
    pregunta_13_data = PreguntasCerradas.objects.values('pregunta_13').annotate(count=Count('pregunta_13')).order_by('pregunta_13')

    # Extraer las opciones de la pregunta_13 y sus frecuencias
    opciones = [item['pregunta_13'] for item in pregunta_13_data]
    frequencies = [item['count'] for item in pregunta_13_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria alto cumplimiento": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria bajo cumplimiento": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('13. Las metas que se proponen en la empresa se cumplen.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que las metas que se proponen se cumplen, lo que indica un buen desempeño organizacional."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe que las metas se cumplan, lo que podría reflejar desafíos en el desempeño organizacional."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre el cumplimiento de las metas en la empresa."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 13.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_13_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_13_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }















def generate_pregunta_14_chart(request):
    # Obtener los datos de la pregunta_14
    pregunta_14_data = PreguntasCerradas.objects.values('pregunta_14').annotate(count=Count('pregunta_14')).order_by('pregunta_14')

    # Extraer las opciones de la pregunta_14 y sus frecuencias
    opciones = [item['pregunta_14'] for item in pregunta_14_data]
    frequencies = [item['count'] for item in pregunta_14_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias para categorías clave
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria de acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('14. Los empleados conocen sus funciones dentro de la empresa.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que los empleados conocen sus funciones, lo que indica una claridad organizacional positiva."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe que los empleados conozcan sus funciones, lo que podría reflejar una falta de claridad en los roles laborales."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre el conocimiento de las funciones por parte de los empleados."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 14.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_14_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_14_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }
















def generate_pregunta_15_chart(request):
    # Obtener los datos de la pregunta_15
    pregunta_15_data = PreguntasCerradas.objects.values('pregunta_15').annotate(count=Count('pregunta_15')).order_by('pregunta_15')

    # Extraer las opciones de la pregunta_15 y sus frecuencias
    opciones = [item['pregunta_15'] for item in pregunta_15_data]
    frequencies = [item['count'] for item in pregunta_15_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_resolucion_positiva = 0
    sumatoria_resolucion_negativa = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_resolucion_positiva += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_resolucion_negativa += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_resolucion_positiva = (sumatoria_resolucion_positiva / total_responses * 100) if total_responses > 0 else 0
    porcentaje_resolucion_negativa = (sumatoria_resolucion_negativa / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria resolución positiva": sumatoria_resolucion_positiva,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria resolución negativa": sumatoria_resolucion_negativa
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('15. En la empresa los problemas entre las personas se resuelven fácilmente.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_resolucion_positiva > sumatoria_resolucion_negativa:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_resolucion_positiva}, {porcentaje_resolucion_positiva:.1f}%) "
            "está de acuerdo en que los problemas se resuelven fácilmente, indicando un buen ambiente de trabajo."
        )
    elif sumatoria_resolucion_negativa > sumatoria_resolucion_positiva:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_resolucion_negativa}, {porcentaje_resolucion_negativa:.1f}%) "
            "no percibe que los problemas se resuelvan fácilmente, lo que podría reflejar tensiones interpersonales."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_resolucion_positiva:.1f}% en acuerdos y {porcentaje_resolucion_negativa:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la resolución de conflictos en la empresa."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 15.\n\n'
        'Incluye sumatorias de respuestas en las categorías de resolución positiva y negativa.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_15_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_15_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }













def generate_pregunta_16_chart(request):
    # Obtener los datos de la pregunta_16
    pregunta_16_data = PreguntasCerradas.objects.values('pregunta_16').annotate(count=Count('pregunta_16')).order_by('pregunta_16')

    # Extraer las opciones de la pregunta_16 y sus frecuencias
    opciones = [item['pregunta_16'] for item in pregunta_16_data]
    frequencies = [item['count'] for item in pregunta_16_data]

    # Calcular el total de respuestas
    total_responses = sum(frequencies)
    
    # Inicializar sumatorias para las categorías
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con las opciones ordenadas
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('16. La forma como está organizada la empresa, es fácil de entender.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "está de acuerdo en que la organización es fácil de entender, lo que indica una buena claridad organizacional."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "no percibe que la organización sea fácil de entender, lo que podría reflejar confusión en las estructuras organizativas."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la claridad organizacional."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 16.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_16_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_16_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }














def generate_pregunta_17_chart(request):
    # Obtener los datos de la pregunta_17
    pregunta_17_data = PreguntasCerradas.objects.values('pregunta_17').annotate(count=Count('pregunta_17')).order_by('pregunta_17')

    # Extraer las opciones de la pregunta_17 y sus frecuencias
    opciones = [item['pregunta_17'] for item in pregunta_17_data]
    frequencies = [item['count'] for item in pregunta_17_data]

    # Calcular el total de respuestas
    total_responses = sum(frequencies)

    # Inicializar sumatorias para las categorías
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con las opciones ordenadas
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('17. Las tareas son supervisadas excesivamente.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "considera que las tareas son supervisadas adecuadamente, lo que indica una buena gestión."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "siente que las tareas son supervisadas en exceso, lo que podría indicar una falta de autonomía."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere una percepción mixta sobre la supervisión de tareas."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 17.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_17_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_17_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }














def generate_pregunta_18_chart(request):
    # Obtener los datos de la pregunta_18
    pregunta_18_data = PreguntasCerradas.objects.values('pregunta_18').annotate(count=Count('pregunta_18')).order_by('pregunta_18')

    # Extraer las opciones de la pregunta_18 y sus frecuencias
    opciones = [item['pregunta_18'] for item in pregunta_18_data]
    frequencies = [item['count'] for item in pregunta_18_data]

    # Calcular el total de respuestas
    total_responses = sum(frequencies)

    # Inicializar sumatorias para las categorías
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con las opciones ordenadas
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('18. En la empresa las relaciones entre las personas son cordiales.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un gran número de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "siente que las relaciones son cordiales, lo que indica un ambiente laboral positivo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "percebe que las relaciones no son cordiales, lo que podría indicar tensiones interpersonales."
        )
    else:
        inferencia = (
            f"Las respuestas están divididas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la cordialidad de las relaciones en la empresa."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 18.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_18_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_18_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }















def generate_pregunta_19_chart(request):
    # Obtener los datos de la pregunta_19
    pregunta_19_data = PreguntasCerradas.objects.values('pregunta_19').annotate(count=Count('pregunta_19')).order_by('pregunta_19')

    # Extraer las opciones de la pregunta_19 y sus frecuencias
    opciones = [item['pregunta_19'] for item in pregunta_19_data]
    frequencies = [item['count'] for item in pregunta_19_data]

    # Calcular el total de respuestas
    total_responses = sum(frequencies)

    # Inicializar sumatorias para las categorías
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Calcular los porcentajes de las sumatorias
    porcentaje_acuerdos = (sumatoria_acuerdos / total_responses * 100) if total_responses > 0 else 0
    porcentaje_desacuerdos = (sumatoria_desacuerdos / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con las opciones ordenadas
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('19. Al interior de la empresa permanentemente hay conflictos.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Un número considerable de encuestados (sumatoria de {sumatoria_acuerdos}, {porcentaje_acuerdos:.1f}%) "
            "siente que hay conflictos permanentes, lo que podría reflejar desafíos en la cultura organizacional."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_desacuerdos}, {porcentaje_desacuerdos:.1f}%) "
            "indica que no perciben conflictos permanentes, sugiriendo un ambiente de trabajo generalmente armonioso."
        )
    else:
        inferencia = (
            f"Las respuestas están equitativamente distribuidas, con {porcentaje_acuerdos:.1f}% en acuerdos y {porcentaje_desacuerdos:.1f}% en desacuerdos, "
            "lo que sugiere opiniones mixtas sobre la existencia de conflictos en el entorno laboral."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 19.\n\n'
        'Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_19_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_19_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }












def generate_pregunta_20_chart(request):
    # Obtener los datos de la pregunta_20
    pregunta_20_data = PreguntasCerradas.objects.values('pregunta_20').annotate(count=Count('pregunta_20')).order_by('pregunta_20')

    # Extraer las opciones de la pregunta_20 y sus frecuencias
    opciones = [item['pregunta_20'] for item in pregunta_20_data]
    frequencies = [item['count'] for item in pregunta_20_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias para categorías de solidaridad
    sumatoria_alto_solidaridad = 0
    sumatoria_bajo_solidaridad = 0

    # Calcular sumatorias basadas en las categorías de solidaridad
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_alto_solidaridad += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_bajo_solidaridad += count

    # Calcular porcentajes de las sumatorias
    porcentaje_alto_solidaridad = (sumatoria_alto_solidaridad / total_responses * 100) if total_responses > 0 else 0
    porcentaje_bajo_solidaridad = (sumatoria_bajo_solidaridad / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria alto solidaridad": sumatoria_alto_solidaridad,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria bajo solidaridad": sumatoria_bajo_solidaridad
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('20. Los empleados son solidarios entre sí.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_alto_solidaridad > sumatoria_bajo_solidaridad:
        inferencia = (
            f"Una cantidad considerable de encuestados (sumatoria de {sumatoria_alto_solidaridad}, {porcentaje_alto_solidaridad:.1f}%) "
            "indica que hay un alto nivel de solidaridad entre empleados, sugiriendo un ambiente de trabajo colaborativo."
        )
    elif sumatoria_bajo_solidaridad > sumatoria_alto_solidaridad:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_bajo_solidaridad}, {porcentaje_bajo_solidaridad:.1f}%) "
            "siente que existe una falta de solidaridad, lo que puede indicar desafíos en el trabajo en equipo."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con {porcentaje_alto_solidaridad:.1f}% en alto solidaridad y {porcentaje_bajo_solidaridad:.1f}% en bajo solidaridad, "
            "lo que indica opiniones divididas sobre la cohesión del equipo."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 20.\n\n'
        'Incluye sumatorias para las categorías de "alto solidaridad" y "bajo solidaridad".\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_20_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_20_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }


















def generate_pregunta_21_chart(request):
    # Obtener los datos de la pregunta_21
    pregunta_21_data = PreguntasCerradas.objects.values('pregunta_21').annotate(count=Count('pregunta_21')).order_by('pregunta_21')

    # Extraer las opciones de la pregunta_21 y sus frecuencias
    opciones = [item['pregunta_21'] for item in pregunta_21_data]
    frequencies = [item['count'] for item in pregunta_21_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias para categorías de tolerancia
    sumatoria_alta_tolerancia = 0
    sumatoria_baja_tolerancia = 0

    # Calcular sumatorias basadas en las categorías de tolerancia
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_alta_tolerancia += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_baja_tolerancia += count

    # Calcular porcentajes de las sumatorias
    porcentaje_alta_tolerancia = (sumatoria_alta_tolerancia / total_responses * 100) if total_responses > 0 else 0
    porcentaje_baja_tolerancia = (sumatoria_baja_tolerancia / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria alta tolerancia": sumatoria_alta_tolerancia,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria baja tolerancia": sumatoria_baja_tolerancia
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('21. Las personas en la empresa son tolerantes.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_alta_tolerancia > sumatoria_baja_tolerancia:
        inferencia = (
            f"Una cantidad considerable de encuestados (sumatoria de {sumatoria_alta_tolerancia}, {porcentaje_alta_tolerancia:.1f}%) "
            "indica que hay un alto nivel de tolerancia entre los empleados, sugiriendo un ambiente laboral positivo."
        )
    elif sumatoria_baja_tolerancia > sumatoria_alta_tolerancia:
        inferencia = (
            f"Una proporción significativa de encuestados (sumatoria de {sumatoria_baja_tolerancia}, {porcentaje_baja_tolerancia:.1f}%) "
            "siente que hay baja tolerancia, lo que podría reflejar tensiones en el equipo de trabajo."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con {porcentaje_alta_tolerancia:.1f}% en alta tolerancia y {porcentaje_baja_tolerancia:.1f}% en baja tolerancia, "
            "lo que sugiere una opinión dividida sobre la tolerancia en la empresa."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 21.\n\n'
        'Incluye sumatorias para las categorías de "alta tolerancia" y "baja tolerancia".\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_21_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_21_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }












def generate_pregunta_22_chart(request):
    # Obtener los datos de la pregunta_22
    pregunta_22_data = PreguntasCerradas.objects.values('pregunta_22').annotate(count=Count('pregunta_22')).order_by('pregunta_22')

    # Extraer las opciones de la pregunta_22 y sus frecuencias
    opciones = [item['pregunta_22'] for item in pregunta_22_data]
    frequencies = [item['count'] for item in pregunta_22_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias para categorías de actualización
    sumatoria_alta_actualizacion = 0
    sumatoria_baja_actualizacion = 0

    # Calcular sumatorias basadas en las categorías de actualización
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_alta_actualizacion += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_baja_actualizacion += count

    # Calcular porcentajes de las sumatorias
    porcentaje_alta_actualizacion = (sumatoria_alta_actualizacion / total_responses * 100) if total_responses > 0 else 0
    porcentaje_baja_actualizacion = (sumatoria_baja_actualizacion / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria alta actualización": sumatoria_alta_actualizacion,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria baja actualización": sumatoria_baja_actualizacion
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('22. Los empleados en la empresa se actualizan en los temas que necesita la organización.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_alta_actualizacion > sumatoria_baja_actualizacion:
        inferencia = (
            f"Una gran proporción de encuestados (sumatoria de {sumatoria_alta_actualizacion}, {porcentaje_alta_actualizacion:.1f}%) "
            "indica que consideran que se actualizan efectivamente en los temas necesarios, sugiriendo un compromiso con el desarrollo profesional."
        )
    elif sumatoria_baja_actualizacion > sumatoria_alta_actualizacion:
        inferencia = (
            f"Una proporción notable de encuestados (sumatoria de {sumatoria_baja_actualizacion}, {porcentaje_baja_actualizacion:.1f}%) "
            "siente que no se están actualizando adecuadamente, lo que podría señalar áreas de mejora en la capacitación del personal."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con {porcentaje_alta_actualizacion:.1f}% en alta actualización y {porcentaje_baja_actualizacion:.1f}% en baja actualización, "
            "lo que sugiere una opinión dividida sobre la efectividad de la formación en la organización."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 22.\n\n'
        'Incluye sumatorias para las categorías de "alta actualización" y "baja actualización".\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_22_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_22_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }















def generate_pregunta_23_chart(request):
    # Obtener los datos de la pregunta_23
    pregunta_23_data = PreguntasCerradas.objects.values('pregunta_23').annotate(count=Count('pregunta_23')).order_by('pregunta_23')

    # Extraer las opciones de la pregunta_23 y sus frecuencias
    opciones = [item['pregunta_23'] for item in pregunta_23_data]
    frequencies = [item['count'] for item in pregunta_23_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias para categorías de autonomía
    sumatoria_alta_autonomia = 0
    sumatoria_baja_autonomia = 0

    # Calcular sumatorias basadas en las categorías de autonomía
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_alta_autonomia += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_baja_autonomia += count

    # Calcular porcentajes de las sumatorias
    porcentaje_alta_autonomia = (sumatoria_alta_autonomia / total_responses * 100) if total_responses > 0 else 0
    porcentaje_baja_autonomia = (sumatoria_baja_autonomia / total_responses * 100) if total_responses > 0 else 0

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria alta autonomía": sumatoria_alta_autonomia,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria baja autonomía": sumatoria_baja_autonomia
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos ordenados
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('23. La empresa apoya la autonomía de sus empleados.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_alta_autonomia > sumatoria_baja_autonomia:
        inferencia = (
            f"Una alta proporción de encuestados (sumatoria de {sumatoria_alta_autonomia}, {porcentaje_alta_autonomia:.1f}%) "
            "indica que la empresa apoya la autonomía de sus empleados, lo cual puede fomentar un ambiente de trabajo positivo."
        )
    elif sumatoria_baja_autonomia > sumatoria_alta_autonomia:
        inferencia = (
            f"Una proporción notable de encuestados (sumatoria de {sumatoria_baja_autonomia}, {porcentaje_baja_autonomia:.1f}%) "
            "siente que la empresa no apoya adecuadamente la autonomía, lo que podría afectar la moral y productividad."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con {porcentaje_alta_autonomia:.1f}% en alta autonomía y {porcentaje_baja_autonomia:.1f}% en baja autonomía, "
            "lo que sugiere una opinión dividida sobre el apoyo a la autonomía en la organización."
        )

    # Descripción al lado derecho de la gráfica incluyendo la inferencia
    description = (
        'Esta gráfica muestra la distribución de respuestas para la pregunta 23.\n\n'
        'Incluye sumatorias para las categorías de "alta autonomía" y "baja autonomía".\n\n'
        f'{inferencia}'
    )

    # Agregar el texto al lado derecho, ajustando el espacio
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    plt.subplots_adjust(right=0.75)

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_23_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_23_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }













def generate_pregunta_24_chart(request):
    # Obtener los datos de la pregunta_24
    pregunta_24_data = PreguntasCerradas.objects.values('pregunta_24').annotate(count=Count('pregunta_24')).order_by('pregunta_24')

    # Validar si hay datos
    if not pregunta_24_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_24 y sus frecuencias
    opciones = [item['pregunta_24'] for item in pregunta_24_data]
    frequencies = [item['count'] for item in pregunta_24_data]

    # Calcular el total de respuestas y porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Categorizar las respuestas
    apoyo = sum(frequencies[i] for i, opcion in enumerate(opciones) if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"])
    falta_apoyo = sum(frequencies[i] for i, opcion in enumerate(opciones) if opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"])

    # Crear un diccionario para ordenar las opciones
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Apoyo al desarrollo (suma)": apoyo,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Falta de apoyo (suma)": falta_apoyo
    }

    # Rellenar las frecuencias en el diccionario
    for opcion, frecuencia in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = frecuencia

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear la figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Crear la gráfica de barras
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=0.5, color='skyblue')

    # Etiquetar las barras con frecuencias y porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate(f'{height} ({percentages_ordenadas[i]:.1f}%)',
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Ajuste vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Personalizar la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('24. La empresa apoya el desarrollo de carrera (ascensos) de sus empleados.')

    # Generar la inferencia
    inferencia = ""
    if apoyo > falta_apoyo:
        inferencia = (
            f"Una proporción significativa de encuestados (suma de {apoyo}, {apoyo / total_responses * 100:.1f}%) "
            "indica que la empresa apoya el desarrollo de carrera de sus empleados, lo que puede fomentar un ambiente positivo."
        )
    elif falta_apoyo > apoyo:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {falta_apoyo}, {falta_apoyo / total_responses * 100:.1f}%) "
            "siente que la empresa no apoya adecuadamente su desarrollo profesional, lo que podría afectar la moral y retención del talento."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con un apoyo y falta de apoyo similares, sugiriendo una opinión dividida acerca del respaldo a los ascensos."
        )

    # Agregar descripción al lado derecho incluyendo la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 24.\n\n"
        "Se incluyen sumatorias para las categorías de 'apoyo' y 'falta de apoyo'.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar la parte derecha para incluir la descripción
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente el diseño
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_24_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_24_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }














def generate_pregunta_25_chart(request):
    # Obtener los datos de la pregunta_25
    pregunta_25_data = PreguntasCerradas.objects.values('pregunta_25').annotate(count=Count('pregunta_25')).order_by('pregunta_25')

    # Validar si hay datos disponibles
    if not pregunta_25_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_25 y sus frecuencias
    opciones = [item['pregunta_25'] for item in pregunta_25_data]
    frequencies = [item['count'] for item in pregunta_25_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Categorización de las respuestas
    apoyo = sum(frequencies[i] for i, opcion in enumerate(opciones) if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"])
    desacuerdo = sum(frequencies[i] for i, opcion in enumerate(opciones) if opcion in ["Totalmente en desacuerdo", "Medianamente en desacuerdo"])
    
    # Crear un diccionario para ordenar las opciones
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Apoyo a la innovación (suma)": apoyo,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Desacuerdo (suma)": desacuerdo
    }

    # Rellenar las frecuencias en el diccionario
    for opcion, frecuencia in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = frecuencia

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear la figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Crear la gráfica de barras
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=0.5, color='skyblue')

    # Etiquetar las barras con las frecuencias y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate(f'{height} ({percentages_ordenadas[i]:.1f}%)',
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Ajuste vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('25. La empresa apoya las sugerencias de los empleados para innovar (en procesos, productos, servicios, etc.).')

    # Generar la inferencia
    inferencia = ""
    if apoyo > desacuerdo:
        inferencia = (
            f"Una cantidad significativa de encuestados (suma de {apoyo}, {apoyo / total_responses * 100:.1f}%) "
            "indica que la empresa apoya las sugerencias de sus empleados para innovar, lo cual podría llevar a mejoras en los procesos y productos."
        )
    elif desacuerdo > apoyo:
        inferencia = (
            f"Una porción notable de encuestados (suma de {desacuerdo}, {desacuerdo / total_responses * 100:.1f}%) "
            "se siente en desacuerdo con el apoyo que la empresa brinda a sus sugerencias, lo que podría indicar áreas de mejora en la cultura de innovación."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, con un apoyo y desacuerdo similares, lo que sugiere una opinión dividida sobre el respaldo a las sugerencias de innovación."
        )

    # Agregar una descripción al lado derecho que incluya la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 25.\n\n"
        "Se incluyen sumatorias para las categorías de 'apoyo' y 'desacuerdo'.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar la parte derecha para incluir la descripción
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_25_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_25_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }

















def generate_pregunta_26_chart(request):
    # Obtener los datos de la pregunta_26
    pregunta_26_data = PreguntasCerradas.objects.values('pregunta_26').annotate(count=Count('pregunta_26')).order_by('pregunta_26')

    # Validar si hay datos disponibles
    if not pregunta_26_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_26 y sus frecuencias
    opciones = [item['pregunta_26'] for item in pregunta_26_data]
    frequencies = [item['count'] for item in pregunta_26_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('26. En la empresa la libertad de expresión se respeta.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que en la empresa la libertad de expresión se respeta, lo que puede contribuir a un entorno laboral positivo y abierto a la comunicación."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "sienten que no se respeta la libertad de expresión, lo que puede señalar áreas de mejora en la cultura organizacional."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, reflejando tanto apoyo como desacuerdo, lo que sugiere una percepción dividida sobre el respeto a la libertad de expresión."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 26.\n\n"
        "Incluye sumatorias de respuestas para las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_26_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_26_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }



















def generate_pregunta_27_chart(request):
    # Obtener los datos de la pregunta_27
    pregunta_27_data = PreguntasCerradas.objects.values('pregunta_27').annotate(count=Count('pregunta_27')).order_by('pregunta_27')

    # Validar si hay datos disponibles
    if not pregunta_27_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_27 y sus frecuencias
    opciones = [item['pregunta_27'] for item in pregunta_27_data]
    frequencies = [item['count'] for item in pregunta_27_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('27. En general, usted se siente bien trabajando en la dependencia actual.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que se sienten bien trabajando en la dependencia actual, lo que podría contribuir a una alta moral y productividad en el lugar de trabajo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "no se sienten bien en su trabajo actual, lo que podría señalar áreas importantes para mejora en la cultura organizacional o condiciones de trabajo."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando una percepción dividida sobre el bienestar en el trabajo, lo que puede requerir atención en la gestión del personal."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 27.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_27_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_27_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }
















def generate_pregunta_28_chart(request):
    # Obtener los datos de la pregunta_28
    pregunta_28_data = PreguntasCerradas.objects.values('pregunta_28').annotate(count=Count('pregunta_28')).order_by('pregunta_28')

    # Validar si hay datos disponibles
    if not pregunta_28_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_28 y sus frecuencias
    opciones = [item['pregunta_28'] for item in pregunta_28_data]
    frequencies = [item['count'] for item in pregunta_28_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('28. Usted se siente bien trabajando en la empresa, en general.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que se sienten bien trabajando en la empresa, lo que sugiere un buen clima laboral y satisfacción general."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "no se sienten bien trabajando en la empresa, lo que podría indicar la necesidad de mejorar el ambiente y las condiciones laborales."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, señalando una percepción mixta sobre el bienestar en el lugar de trabajo, lo que puede requerir atención por parte de la gestión."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 28.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_28_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_28_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }



















def generate_pregunta_29_chart(request):
    # Obtener los datos de la pregunta_29
    pregunta_29_data = PreguntasCerradas.objects.values('pregunta_29').annotate(count=Count('pregunta_29')).order_by('pregunta_29')

    # Validar si hay datos disponibles
    if not pregunta_29_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_29 y sus frecuencias
    opciones = [item['pregunta_29'] for item in pregunta_29_data]
    frequencies = [item['count'] for item in pregunta_29_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('29. En general, la empresa paga los salarios que cada quien se merece.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "cree que la empresa paga salarios justos, lo que sugiere satisfacción respecto a la compensación laboral."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "manifiesta indiferencia o desacuerdo sobre la equidad salarial, lo que podría señalar una necesidad de revisar las políticas de compensación."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando una percepción mixta sobre la justicia en la compensación salarial, "
            "lo cual podría requerir una investigación más profunda para abordar inquietudes existentes."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 29.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_29_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_29_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }
















def generate_pregunta_30_chart(request):
    # Obtener los datos de la pregunta_30
    pregunta_30_data = PreguntasCerradas.objects.values('pregunta_30').annotate(count=Count('pregunta_30')).order_by('pregunta_30')

    # Validar si hay datos disponibles
    if not pregunta_30_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_30 y sus frecuencias
    opciones = [item['pregunta_30'] for item in pregunta_30_data]
    frequencies = [item['count'] for item in pregunta_30_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('30. Frente a entidades parecidas, la empresa es fuerte.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "confirma que la empresa es fuerte en comparación con entidades similares, lo que podría reflejar una buena imagen de marca y posicionamiento en el mercado."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "no creen que la empresa sea fuerte frente a sus competidores, lo que podría indicar desconfianza en la estrategia empresarial o en la competitividad."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando una percepción mixta hacia la fuerza de la empresa frente a entidades similares, "
            "lo que podría requerir un análisis más profundo para entender las inquietudes entre los empleados."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 30.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_30_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_30_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }


















def generate_pregunta_31_chart(request):
    # Obtener los datos de la pregunta_31
    pregunta_31_data = PreguntasCerradas.objects.values('pregunta_31').annotate(count=Count('pregunta_31')).order_by('pregunta_31')

    # Validar si hay datos disponibles
    if not pregunta_31_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_31 y sus frecuencias
    opciones = [item['pregunta_31'] for item in pregunta_31_data]
    frequencies = [item['count'] for item in pregunta_31_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('31. Esta organización le cumple a sus clientes.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "cree que la organización cumple con sus promesas a los clientes, lo que sugiere una fuerte confianza en la atención y el servicio al cliente."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad notable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "no consideran que la organización cumple a sus clientes, indicando posibles áreas de mejora en el servicio o la entrega de productos."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando una percepción mixta sobre la capacidad de la organización para cumplir con sus clientes, "
            "lo que podría requerir una evaluación más detallada para identificar inquietudes individuales."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 31.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_31_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_31_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }















def generate_pregunta_32_chart(request):
    # Obtener los datos de la pregunta_32
    pregunta_32_data = PreguntasCerradas.objects.values('pregunta_32').annotate(count=Count('pregunta_32')).order_by('pregunta_32')

    # Validar si hay datos disponibles
    if not pregunta_32_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_32 y sus frecuencias
    opciones = [item['pregunta_32'] for item in pregunta_32_data]
    frequencies = [item['count'] for item in pregunta_32_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('32. Si usted recibiera una oferta de trabajo de otra organización se iría, siendo las condiciones de la otra las mismas.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que estarían dispuestos a cambiarse a otra organización bajo las mismas condiciones, sugiriendo un alto deseo de mejorar sus oportunidades laborales."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad considerable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "prefiere quedarse, lo que podría indicar satisfacción con la organización actual y una percepción de que las oportunidades son mejores en su puesto actual."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, mostrando una mezcla de opiniones sobre si cambiarían a otra organización bajo las mismas condiciones, "
            "lo que podría requerir una evaluación más detallada sobre la percepción del ambiente laboral y las oportunidades ofrecidas."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 32.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_32_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_32_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }



















def generate_pregunta_33_chart(request):
    # Obtener los datos de la pregunta_33
    pregunta_33_data = PreguntasCerradas.objects.values('pregunta_33').annotate(count=Count('pregunta_33')).order_by('pregunta_33')

    # Validar si hay datos disponibles
    if not pregunta_33_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_33 y sus frecuencias
    opciones = [item['pregunta_33'] for item in pregunta_33_data]
    frequencies = [item['count'] for item in pregunta_33_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('33. Si usted recibiera una oferta de trabajo de otra organización se iría, siendo las condiciones de la otra mucho mejores.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que estarían dispuestos a aceptar una oferta de trabajo de otra organización si las condiciones son considerablemente mejores."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad considerable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "prefiere seguir en su trabajo actual, lo que podría sugerir satisfacción con su puesto actual o lealtad hacia la organización."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando que los encuestados presentan sentimientos mixtos con respecto a cambiarse a otra organización, "
            "lo que podría indicar una necesidad de evaluar más a fondo los factores que motivarían su decisión."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 33.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_33_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_33_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }



















def generate_pregunta_34_chart(request):
    # Obtener los datos de la pregunta_34
    pregunta_34_data = PreguntasCerradas.objects.values('pregunta_34').annotate(count=Count('pregunta_34')).order_by('pregunta_34')

    # Validar si hay datos disponibles
    if not pregunta_34_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_34 y sus frecuencias
    opciones = [item['pregunta_34'] for item in pregunta_34_data]
    frequencies = [item['count'] for item in pregunta_34_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('34. Las condiciones de su sitio de trabajo son adecuadas para desempeñarse bien.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "considera que las condiciones de su sitio de trabajo son adecuadas para desempeñarse bien, lo que sugiere un ambiente laboral favorable."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad significativa de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "indica que las condiciones laborales no son satisfactorias, lo que podría señalar áreas de mejora necesarias para el bienestar de los empleados."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando sentimientos mixtos sobre la adecuación de las condiciones de trabajo, "
            "lo que podría requerir una investigación adicional para comprender mejor las preocupaciones y necesidades del personal."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 34.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_34_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_34_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }












def generate_pregunta_35_chart(request):
    # Obtener los datos de la pregunta_35
    pregunta_35_data = PreguntasCerradas.objects.values('pregunta_35').annotate(count=Count('pregunta_35')).order_by('pregunta_35')

    # Validar si hay datos disponibles
    if not pregunta_35_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_35 y sus frecuencias
    opciones = [item['pregunta_35'] for item in pregunta_35_data]
    frequencies = [item['count'] for item in pregunta_35_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('35. Es notable la presencia de grupos cerrados en los cuales se refugian sus integrantes.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que perciben la presencia de grupos cerrados en su entorno, lo que sugiere una tendencia hacia la cohesión grupal."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad considerable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "se opone a la idea de que existen grupos cerrados, lo que podría indicar una percepción de apertura y diversidad en el ambiente laboral."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando sentimientos mixtos sobre la existencia de grupos cerrados, "
            "lo que podría requerir un análisis adicional para entender las dinámicas del grupo y las relaciones entre compañeros."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 35.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo, para analizar la percepción de cohesión grupal.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_35_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_35_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }




















def generate_pregunta_36_chart(request):
    # Obtener los datos de la pregunta_36
    pregunta_36_data = PreguntasCerradas.objects.values('pregunta_36').annotate(count=Count('pregunta_36')).order_by('pregunta_36')

    # Validar si hay datos disponibles
    if not pregunta_36_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_36 y sus frecuencias
    opciones = [item['pregunta_36'] for item in pregunta_36_data]
    frequencies = [item['count'] for item in pregunta_36_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular sumatorias para categorías específicas
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Definir un orden fijo para las categorías de respuesta
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias reales
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer datos en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('36. La cantidad de tareas que tiene su cargo es mayor a la de otros cargos que se le parecen.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "siente que la cantidad de tareas que tienen a su cargo es mayor en comparación con otros puestos similares, lo que puede reflejar una carga de trabajo significativa."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una cantidad considerable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "indica que no perciben una carga de trabajo superior en relación a otros puestos similares, lo que podría sugerir una percepción de distribución de tareas justa."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, indicando que hay sentimientos mixtos sobre la carga de trabajo en relación a otros puestos, "
            "lo que podría requerir un análisis más profundo de las expectativas laborales y la carga asignada."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 36.\n\n"
        "Incluye sumatorias de las respuestas agrupadas en categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_36_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_36_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }























def generate_pregunta_37_chart(request):
    # Obtener los datos de la pregunta_37
    pregunta_37_data = PreguntasCerradas.objects.values('pregunta_37').annotate(count=Count('pregunta_37')).order_by('pregunta_37')

    # Validar si hay datos disponibles
    if not pregunta_37_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_37 y sus frecuencias
    opciones = [item['pregunta_37'] for item in pregunta_37_data]
    frequencies = [item['count'] for item in pregunta_37_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('37. Su jefe sabe cómo hacer el trabajo de sus subalternos.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "afirma que su jefe sabe cómo hacer el trabajo de sus subalternos, lo que indica un ambiente laboral donde se valora el conocimiento y la competencia del liderazgo."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Una parte considerable de los encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "considera que su jefe no está familiarizado con las tareas que realizan los subalternos, lo que podría evidenciar una falta de entrenamiento o preparación en la alta dirección."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, lo que indica diversas percepciones sobre la competencia de liderazgo, "
            "lo que puede llevar a un análisis más profundo de las dinámicas de equipo y la formación de líderes."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 37.\n\n"
        "Incluye sumatorias de las respuestas agrupadas en categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_37_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_37_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }

















def generate_pregunta_38_chart(request):
    # Obtener los datos de la pregunta_38
    pregunta_38_data = PreguntasCerradas.objects.values('pregunta_38').annotate(count=Count('pregunta_38')).order_by('pregunta_38')

    # Validar si hay datos disponibles
    if not pregunta_38_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_38 y sus frecuencias
    opciones = [item['pregunta_38'] for item in pregunta_38_data]
    frequencies = [item['count'] for item in pregunta_38_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('38. Su jefe sabe cómo premiar a sus subalternos.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que su jefe sabe cómo premiar a sus subalternos, lo que puede reflejar una buena práctica de reconocimiento y motivación en el ambiente laboral."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número considerable de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "considera que su jefe no tiene claro cómo premiar a sus subalternos, lo que sugiere que podría haber oportunidades para mejorar la cultura de reconocimiento."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, lo que sugiere que hay opiniones mixtas sobre la capacidad de liderazgo en términos de reconocimiento de logros, "
            "lo que podría indicar la necesidad de que las empresas trabajen en su cultura de recompensa."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 38.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_38_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_38_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }













def generate_pregunta_39_chart(request):
    # Obtener los datos de la pregunta_39
    pregunta_39_data = PreguntasCerradas.objects.values('pregunta_39').annotate(count=Count('pregunta_39')).order_by('pregunta_39')

    # Validar si hay datos disponibles
    if not pregunta_39_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_39 y sus frecuencias
    opciones = [item['pregunta_39'] for item in pregunta_39_data]
    frequencies = [item['count'] for item in pregunta_39_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('39. Su jefe sabe cómo sancionar a sus subalternos.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que su jefe sabe cómo sancionar a sus subalternos, lo que podría sugerir un enfoque efectivo y equitativo en la gestión de personal."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número significativo de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "considera que su jefe no sabe cómo sancionar adecuadamente a sus subalternos, lo que podría reflejar una falta de claridad en las políticas disciplinarias."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, lo que sugiere que hay opiniones divididas respecto a la capacidad de liderazgo en términos de sanciones, "
            "lo que podría indicar la necesidad de revisar la comunicación y aplicación de las sanciones en el ambiente laboral."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 39.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_39_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_39_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }


















def generate_pregunta_40_chart(request):
    # Obtener los datos de la pregunta_40
    pregunta_40_data = PreguntasCerradas.objects.values('pregunta_40').annotate(count=Count('pregunta_40')).order_by('pregunta_40')

    # Validar si hay datos disponibles
    if not pregunta_40_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_40 y sus frecuencias
    opciones = [item['pregunta_40'] for item in pregunta_40_data]
    frequencies = [item['count'] for item in pregunta_40_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('40. Su puesto de trabajo tiene variedad en la forma de ejecutar las tareas.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que en su puesto de trabajo hay variedad en la forma de ejecutar las tareas, lo que sugiere un ambiente laboral flexible y adaptable."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número significativo de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "siente que su puesto de trabajo carece de diversidad en los métodos de ejecución, lo que podría indicar un entorno rígido que no fomenta la creatividad."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, lo que sugiere que la percepción sobre la diversidad en las tareas es variada, "
            "indicando una posible necesidad de revisar las políticas de trabajo para fomentar un entorno más dinámico."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 40.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_40_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_40_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }



















def generate_pregunta_41_chart(request):
    # Obtener los datos de la pregunta_41
    pregunta_41_data = PreguntasCerradas.objects.values('pregunta_41').annotate(count=Count('pregunta_41')).order_by('pregunta_41')

    # Validar si hay datos disponibles
    if not pregunta_41_data:
        return "No hay datos disponibles para generar la gráfica."

    # Extraer las opciones de la pregunta_41 y sus frecuencias
    opciones = [item['pregunta_41'] for item in pregunta_41_data]
    frequencies = [item['count'] for item in pregunta_41_data]

    # Calcular el total de respuestas y los porcentajes
    total_responses = sum(frequencies)
    percentages = [(count / total_responses * 100) if total_responses > 0 else 0 for count in frequencies]

    # Inicializar variables para las sumatorias
    sumatoria_acuerdos = 0
    sumatoria_desacuerdos = 0

    # Calcular las sumatorias basadas en las categorías
    for opcion, count in zip(opciones, frequencies):
        if opcion in ["Totalmente de acuerdo", "Medianamente de acuerdo"]:
            sumatoria_acuerdos += count
        elif opcion in ["Medianamente en desacuerdo", "Totalmente en desacuerdo"]:
            sumatoria_desacuerdos += count

    # Crear un diccionario con las categorías en el orden deseado
    categorias_ordenadas = {
        "Totalmente de acuerdo": 0,
        "Medianamente de acuerdo": 0,
        "Sumatoria acuerdos": sumatoria_acuerdos,
        "Medianamente en desacuerdo": 0,
        "Totalmente en desacuerdo": 0,
        "Sumatoria de desacuerdos": sumatoria_desacuerdos
    }

    # Rellenar el diccionario con las frecuencias de las opciones obtenidas
    for opcion, count in zip(opciones, frequencies):
        if opcion in categorias_ordenadas:
            categorias_ordenadas[opcion] = count

    # Extraer las opciones y frecuencias en el orden deseado
    opciones_ordenadas = list(categorias_ordenadas.keys())
    frequencies_ordenadas = list(categorias_ordenadas.values())
    percentages_ordenadas = [(freq / total_responses * 100) if total_responses > 0 else 0 for freq in frequencies_ordenadas]

    # Crear una nueva figura para la gráfica
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras con los datos en el orden deseado
    bars = ax.bar(opciones_ordenadas, frequencies_ordenadas, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('41. Usted encuentra congruencia entre lo que busca en su vida laboral y lo que le ofrece su puesto de trabajo.')

    # Mostrar las frecuencias y porcentajes encima de las barras
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages_ordenadas[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Generar la inferencia
    inferencia = ""
    if sumatoria_acuerdos > sumatoria_desacuerdos:
        inferencia = (
            f"Una mayoría de los encuestados (suma de {sumatoria_acuerdos}, {sumatoria_acuerdos / total_responses * 100:.1f}%) "
            "indica que encuentra congruencia entre sus expectativas laborales y lo que su puesto ofrece, sugiriendo un alto nivel de satisfacción y alineación."
        )
    elif sumatoria_desacuerdos > sumatoria_acuerdos:
        inferencia = (
            f"Un número significativo de encuestados (suma de {sumatoria_desacuerdos}, {sumatoria_desacuerdos / total_responses * 100:.1f}%) "
            "siente que hay una discrepancia entre lo que busca y lo que recibe en su puesto, lo que puede señalar una insatisfacción laboral."
        )
    else:
        inferencia = (
            f"Las respuestas están equilibradas, lo que sugiere una percepción mixta respecto a la congruencia laboral, "
            "indicando oportunidades para que la organización escuche y ajuste sus ofertas laborales."
        )

    # Descripción al lado derecho de la gráfica que incluye la inferencia
    description = (
        "Esta gráfica muestra la distribución de respuestas para la pregunta 41.\n\n"
        "Incluye sumatorias de respuestas en las categorías de acuerdo y desacuerdo.\n\n"
        f"{inferencia}"
    )
    
    # Ajustar el texto a la derecha
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=10,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=10))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_41_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_41_chart.png')

    return {
        'chart_url': chart_url,
        'percentages_ordenadas': percentages_ordenadas,
        'inferencia': inferencia
    }













# Obtener los datos de la pregunta_42
def procesar_respuestas3(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    respuestas_situacion_1 = PreguntaAbierta.objects.values_list('pregunta_42_situacion_1', flat=True)
    respuestas_situacion_2 = PreguntaAbierta.objects.values_list('pregunta_42_situacion_2', flat=True)
    respuestas = list(respuestas_situacion_1) + list(respuestas_situacion_2)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias, siempre 10, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              'situaciones, anécdotas, historias internas o algo típico, que refleje lo que distingue la cultura de esta organización frente a las que se le parecen. Algo que permita decir: "esto solo pasa aquí".'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", ame un conteo y un porcentaje para cada categoria.'
              'los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , Siempre se va a usar ese formato, siempre.')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('Situaciones, anécdotas, historias internas o algo típico')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_42_chart_bar3.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_42_chart_bar3.png')

    return chart_url

















def generate_pregunta_43_opcion_1_chart(request):
    # Obtener los datos de la pregunta 43, opción 1
    pregunta_43_opcion_1_data = PreguntaImportancia.objects.values('pregunta_43_opcion_1').annotate(count=Count('pregunta_43_opcion_1')).order_by('pregunta_43_opcion_1')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_43_opcion_1'] for item in pregunta_43_opcion_1_data]
    frequencies = [item['count'] for item in pregunta_43_opcion_1_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Importancia de la razón por la cual trabaja aquí: "Me siento bien con mis compañeros"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 43, opción 1.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_43_opcion_1_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_43_opcion_1_chart.png')

    return chart_url










def generate_pregunta_43_opcion_2_chart(request):
    # Obtener los datos de la pregunta 43, opción 2
    pregunta_43_opcion_2_data = PreguntaImportancia.objects.values('pregunta_43_opcion_2').annotate(count=Count('pregunta_43_opcion_2')).order_by('pregunta_43_opcion_2')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_43_opcion_2'] for item in pregunta_43_opcion_2_data]
    frequencies = [item['count'] for item in pregunta_43_opcion_2_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Importancia de la razón por la cual trabaja aquí: "Puedo ayudar a organizar los equipos de trabajo"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 43, opción 2.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_43_opcion_2_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_43_opcion_2_chart.png')

    return chart_url








def generate_pregunta_43_opcion_3_chart(request):
    # Obtener los datos de la pregunta 43, opción 3
    pregunta_43_opcion_3_data = PreguntaImportancia.objects.values('pregunta_43_opcion_3').annotate(count=Count('pregunta_43_opcion_3')).order_by('pregunta_43_opcion_3')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_43_opcion_3'] for item in pregunta_43_opcion_3_data]
    frequencies = [item['count'] for item in pregunta_43_opcion_3_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Importancia de la razón por la cual trabaja aquí: "Puedo avanzar hacia las metas que me he propuesto en la vida"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 43, opción 3.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_43_opcion_3_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_43_opcion_3_chart.png')

    return chart_url
    



def generate_pregunta_43_chart(request):
    # Obtener los datos de la pregunta 43 para las tres opciones
    pregunta_43_opcion_1_data = PreguntaImportancia.objects.values('pregunta_43_opcion_1').annotate(count=Count('pregunta_43_opcion_1')).order_by('pregunta_43_opcion_1')
    pregunta_43_opcion_2_data = PreguntaImportancia.objects.values('pregunta_43_opcion_2').annotate(count=Count('pregunta_43_opcion_2')).order_by('pregunta_43_opcion_2')
    pregunta_43_opcion_3_data = PreguntaImportancia.objects.values('pregunta_43_opcion_3').annotate(count=Count('pregunta_43_opcion_3')).order_by('pregunta_43_opcion_3')

    # Extraer las respuestas y sus frecuencias para cada opción
    opciones_1 = [item['pregunta_43_opcion_1'] for item in pregunta_43_opcion_1_data]
    opciones_2 = [item['pregunta_43_opcion_2'] for item in pregunta_43_opcion_2_data]
    opciones_3 = [item['pregunta_43_opcion_3'] for item in pregunta_43_opcion_3_data]
    frequencies_1 = [item['count'] for item in pregunta_43_opcion_1_data]
    frequencies_2 = [item['count'] for item in pregunta_43_opcion_2_data]
    frequencies_3 = [item['count'] for item in pregunta_43_opcion_3_data]

    # Calcular porcentajes para cada opción
    total_responses_1 = sum(frequencies_1)
    total_responses_2 = sum(frequencies_2)
    total_responses_3 = sum(frequencies_3)
    percentages_1 = [count / total_responses_1 * 100 for count in frequencies_1]
    percentages_2 = [count / total_responses_2 * 100 for count in frequencies_2]
    percentages_3 = [count / total_responses_3 * 100 for count in frequencies_3]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.2

    # Crear la gráfica de barras
    x = range(len(opciones_1))
    ax.bar(x, frequencies_1, width=bar_width, label='Me siento bien con mis compañeros', color='#1f77b4')
    ax.bar([i + bar_width for i in x], frequencies_2, width=bar_width, label='Puedo ayudar a organizar los equipos de trabajo', color='#ff7f0e')
    ax.bar([i + bar_width * 2 for i in x], frequencies_3, width=bar_width, label='Puedo avanzar hacia las metas que me he propuesto en la vida', color='#808080')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Respuestas a la pregunta 43')
    ax.set_xticks([i + bar_width for i in x])
    ax.set_xticklabels(['Relaciones Interpersonales', 'Liderazgo', 'Logros Personales'])
    ax.legend()

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(ax.patches):
        height = rect.get_height()
        if i < len(opciones_1):
            ax.annotate('{} ({:.1f}%)'.format(height, percentages_1[i]),
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # Desplazamiento vertical del texto
                        textcoords="offset points",
                        ha='center', va='bottom')
        elif i < len(opciones_1) * 2:
            ax.annotate('{} ({:.1f}%)'.format(height, percentages_2[i - len(opciones_1)]),
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # Desplazamiento vertical del texto
                        textcoords="offset points",
                        ha='center', va='bottom')
        else:
            ax.annotate('{} ({:.1f}%)'.format(height, percentages_3[i - len(opciones_1) * 2]),
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # Desplazamiento vertical del texto
                        textcoords="offset points",
                        ha='center', va='bottom')

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_43_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_43_chart.png')

    return chart_url










def generate_pregunta_44_opcion_1_chart(request):
    # Obtener los datos de la pregunta 44, opción 1
    pregunta_44_opcion_1_data = PreguntaImportancia.objects.values('pregunta_44_opcion_1').annotate(count=Count('pregunta_44_opcion_1')).order_by('pregunta_44_opcion_1')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_44_opcion_1'] for item in pregunta_44_opcion_1_data]
    frequencies = [item['count'] for item in pregunta_44_opcion_1_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Fuentes de poder tiene mayor influencia en esta entidad: "Las directivas"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 44, opción 1.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_44_opcion_1_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_44_opcion_1_chart.png')

    return chart_url









def generate_pregunta_44_opcion_2_chart(request):
    # Obtener los datos de la pregunta 44, opción 2
    pregunta_44_opcion_2_data = PreguntaImportancia.objects.values('pregunta_44_opcion_2').annotate(count=Count('pregunta_44_opcion_2')).order_by('pregunta_44_opcion_2')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_44_opcion_2'] for item in pregunta_44_opcion_2_data]
    frequencies = [item['count'] for item in pregunta_44_opcion_2_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Fuentes de poder tiene mayor influencia en esta entidad: "Los empleados"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 44, opción 2.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_44_opcion_2_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_44_opcion_2_chart.png')

    return chart_url









def generate_pregunta_44_opcion_3_chart(request):
    # Obtener los datos de la pregunta 44, opción 3
    pregunta_44_opcion_3_data = PreguntaImportancia.objects.values('pregunta_44_opcion_3').annotate(count=Count('pregunta_44_opcion_3')).order_by('pregunta_44_opcion_3')

    # Extraer las respuestas y sus frecuencias
    opciones = [item['pregunta_44_opcion_3'] for item in pregunta_44_opcion_3_data]
    frequencies = [item['count'] for item in pregunta_44_opcion_3_data]

    # Calcular porcentajes
    total_responses = sum(frequencies)
    percentages = [count / total_responses * 100 for count in frequencies]

    # Crear una nueva figura
    fig, ax = plt.subplots(figsize=(16, 6))

    # Ajustar el ancho de las barras
    bar_width = 0.5

    # Crear la gráfica de barras
    bars = ax.bar(opciones, frequencies, width=bar_width, color='skyblue')

    # Personalizar la apariencia de la gráfica
    ax.set_xlabel('Opciones')
    ax.set_ylabel('Frecuencia')
    ax.set_title('Fuentes de poder tiene mayor influencia en esta entidad: "Factores externos a la entidad"')

    # Mostrar la cantidad exacta de veces que se ha respondido cada opción en el eje y y los porcentajes
    for i, rect in enumerate(bars):
        height = rect.get_height()
        ax.annotate('{} ({:.1f}%)'.format(height, percentages[i]),
                    xy=(rect.get_x() + rect.get_width() / 2, height),
                    xytext=(0, 3),  # Desplazamiento vertical del texto
                    textcoords="offset points",
                    ha='center', va='bottom')

    # Descripción al lado derecho de la gráfica
    description = 'Esta gráfica muestra la distribución de respuestas para la pregunta 44, opción 3.\n\n'
    description += 'Los porcentajes indican la proporción de respuestas en relación con el total de respuestas.\n\n'
    description += '1) La menos importante\n2) Medianamente importante\n3) La más importante'
    ax.text(1.05, 0.5, description, transform=ax.transAxes, fontsize=12,
            va='center', ha='left', wrap=True, bbox=dict(facecolor='none', edgecolor='black', pad=20))

    # Ajustar automáticamente los parámetros de la figura
    fig.tight_layout()

    # Guardar la imagen en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_44_opcion_3_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    # Obtener la URL de la imagen
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_44_opcion_3_chart.png')

    return chart_url

















# Obtener los datos de la pregunta_44 abierta 
def generate_pregunta_44_opcion_4_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    pregunta_44_opcion_4 = PreguntaImportancia.objects.values_list('pregunta_44_opcion_4', flat=True)
    respuestas = list(pregunta_44_opcion_4) + list(pregunta_44_opcion_4)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 3 categorias, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              'Fuentes de poder con influencia en esta entidad y que son Factores externos a la entidad.'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'Usa solo los textos para crear las categorias, el contexto es solo para darte una idea del tema al cual deben hacer referencia las categorias'
              'para crear las categorias, usa terminos en relacion con una jerga profesional en la psicologia organizacional'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , SIEMPRE se va a usar ese formato sin alteraciones, siempre.'
              'no olvide el formato de entrega')

    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)


    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('Textos por Categoría')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_44_opcion_4_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_44_opcion_4_chart.png')

    return chart_url








def generate_pregunta_44_chart(request):
    # Definir las categorías y los niveles de influencia
    categorias = ["Directivas", "Empleados", "Factores externos a la entidad"]
    niveles_influencia = ["Menor influencia", "Mediana influencia", "Mayor influencia"]

    # Inicialización de conteos de influencias para cada categoría
    conteos = {
        "Directivas": [0, 0, 0],  # [Menor influencia, Mediana influencia, Mayor influencia]
        "Empleados": [0, 0, 0],
        "Factores externos a la entidad": [0, 0, 0]
    }

    # Obtener los datos y contar respuestas por nivel de influencia para cada categoría
    datos_opciones = {
        "Directivas": "pregunta_44_opcion_1",
        "Empleados": "pregunta_44_opcion_2",
        "Factores externos a la entidad": "pregunta_44_opcion_3"
    }

    for categoria, campo in datos_opciones.items():
        datos = PreguntaImportancia.objects.values(campo).annotate(count=Count(campo))
        
        for item in datos:
            valor = item[campo]
            if valor in [1, 2, 3]:  # Mapear 1 -> Menor, 2 -> Mediana, 3 -> Mayor
                conteos[categoria][valor - 1] = item['count']

    # Calcular porcentajes y conteos para cada nivel de influencia
    porcentajes = {categoria: [] for categoria in categorias}
    conteos_totales = {categoria: [] for categoria in categorias}

    for categoria in categorias:
        total_respuestas = sum(conteos[categoria])
        for idx in range(3):
            porcentaje = (conteos[categoria][idx] / total_respuestas * 100) if total_respuestas > 0 else 0
            porcentajes[categoria].append(porcentaje)
            conteos_totales[categoria].append(conteos[categoria][idx])

    # Preparar los datos para la gráfica
    menor_influencia = [porcentajes["Directivas"][0], porcentajes["Empleados"][0], porcentajes["Factores externos a la entidad"][0]]
    mediana_influencia = [porcentajes["Directivas"][1], porcentajes["Empleados"][1], porcentajes["Factores externos a la entidad"][1]]
    mayor_influencia = [porcentajes["Directivas"][2], porcentajes["Empleados"][2], porcentajes["Factores externos a la entidad"][2]]

    menor_conteo = [conteos_totales["Directivas"][0], conteos_totales["Empleados"][0], conteos_totales["Factores externos a la entidad"][0]]
    mediana_conteo = [conteos_totales["Directivas"][1], conteos_totales["Empleados"][1], conteos_totales["Factores externos a la entidad"][1]]
    mayor_conteo = [conteos_totales["Directivas"][2], conteos_totales["Empleados"][2], conteos_totales["Factores externos a la entidad"][2]]

    # Crear el gráfico de barras
    fig, ax = plt.subplots(figsize=(10, 6))
    bar_width = 0.2
    x = np.arange(len(categorias))  # Posiciones de las categorías principales

    # Graficar barras para cada nivel de influencia dentro de cada localización
    bars_menor = ax.bar(x - bar_width, menor_influencia, width=bar_width, color='#1f77b4', label='Menor influencia')
    bars_mediana = ax.bar(x, mediana_influencia, width=bar_width, color='#ff7f0e', label='Mediana influencia')
    bars_mayor = ax.bar(x + bar_width, mayor_influencia, width=bar_width, color='#7f7f7f', label='Mayor influencia')

    # Etiquetas y título
    ax.set_xlabel('Localización')
    ax.set_ylabel('Porcentaje (%)')
    ax.set_title('Distribución de influencia según la localización del poder')
    ax.set_xticks(x)
    ax.set_xticklabels(categorias)
    ax.legend(title='Nivel de Influencia')

    # Anotaciones para porcentaje y conteo en cada barra
    for bars, counts, percentages in zip([bars_menor, bars_mediana, bars_mayor], [menor_conteo, mediana_conteo, mayor_conteo], [menor_influencia, mediana_influencia, mayor_influencia]):
        for bar, count, percentage in zip(bars, counts, percentages):
            height = bar.get_height()
            if height > 0:
                ax.annotate(f'{percentage:.2f}%\n({count})',
                            xy=(bar.get_x() + bar.get_width() / 2, height),
                            xytext=(0, 3), textcoords="offset points",
                            ha='center', va='bottom')

    fig.tight_layout()

    # Guardar la imagen
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'pregunta_44_chart.png')
    fig.savefig(chart_path)
    plt.close(fig)  # Cierra la figura y libera memoria

    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_44_chart.png')
    return chart_url



































# Obtener los datos de la pregunta_45
def generate_pregunta_45_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    defecto_1 = PreguntaAbiertaDefectos.objects.values_list('defecto_1', flat=True)
    defecto_2 = PreguntaAbiertaDefectos.objects.values_list('defecto_2', flat=True)
    defecto_3 = PreguntaAbiertaDefectos.objects.values_list('defecto_3', flat=True)
    respuestas = list(defecto_1) + list(defecto_2) + list(defecto_3) 

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = (
        "A partir de los textos suministrados (cada uno separado por el carácter '|'), genera exactamente 10 categorías de defectos observados en una entidad. Para ello, cumple con las siguientes condiciones:\n"
        "1. Los títulos de las categorías deben relacionarse directamente con conceptos organizacionales y emplear la terminología propia de la psicología organizacional.\n"
        "2. Utiliza exclusivamente aquellos textos que aporten información significativa, omitiendo expresiones vacías o irrelevantes tales como 'Ninguna', 'N/A', '...', etc.\n"
        "3. Realiza un conteo de la frecuencia con que aparece cada categoría en los textos y calcula el porcentaje correspondiente en función del total de textos procesados.\n"
        "4. Excluye de la salida cualquier categoría cuyo porcentaje resulte en 0%, para evitar información innecesaria.\n"
        "5. Presenta la salida de manera estricta en el siguiente formato de tabla, sin incluir información adicional:\n"
        "   | Categoria | Conteo | Porcentaje |\n"
    )



    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('45. Defectos de la entidad')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_45_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_45_chart.png')

    return chart_url










# Obtener los datos de la pregunta_46
def generate_pregunta_46_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    virtud_1 = PreguntaAbiertaVirtudes.objects.values_list('virtud_1', flat=True)
    virtud_2 = PreguntaAbiertaVirtudes.objects.values_list('virtud_2', flat=True)
    virtud_3 = PreguntaAbiertaVirtudes.objects.values_list('virtud_3', flat=True)
    respuestas = list(virtud_1) + list(virtud_2) + list(virtud_3) 

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias, siempre 10, de almenos dos palabras, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              'Mencione virtudes de esta entidad.'
              'para crear las categorias, usa los terminos en relacion con una jerga profesional en la psicologia organizacional'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias. No cree una categoria llamada "otros"'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | . SIEMPRE se va a usar ese formato sin alteraciones, siempre.'
              'no es necesario entregar más informacion ademas de las categorias, el conteo y el porcentaje en  el formato solicitado.')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('46. Virtudes de la entidad')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_46_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_46_chart.png')

    return chart_url











# Obtener los datos de la pregunta_47
def generate_pregunta_47_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    habito_1 = PreguntaAbiertaHabitos.objects.values_list('habito_1', flat=True)
    habito_2 = PreguntaAbiertaHabitos.objects.values_list('habito_2', flat=True)
    respuestas = list(habito_1) + list(habito_2)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = (
        "A partir de los textos suministrados (cada uno separado por el carácter '|'), genera exactamente 10 categorías que describan hábitos diarios que se practican en la empresa para mejorar los resultados. Sigue estas directrices:\n"
        "1. Los títulos de las categorías deben estar relacionados con conceptos organizacionales y utilizar terminología propia de la psicología organizacional.\n"
        "2. Emplea únicamente los textos que aporten información relevante, omitiendo aquellos que solo contengan expresiones vacías o irrelevantes como 'Ninguna', 'N/A', '...', etc.\n"
        "3. Asegúrate de que los títulos de las categorías estén compuestos por más de una palabra, de modo que se describa con precisión el tema correspondiente.\n"
        "4. Realiza un conteo de la frecuencia con la que cada categoría es mencionada en los textos y calcula el porcentaje en función del total de textos procesados.\n"
        "5. Excluye de la salida cualquier categoría cuyo porcentaje sea 0%, evitando información innecesaria.\n"
        "6. La salida debe presentarse de forma estricta en el siguiente formato, sin incluir información adicional:\n"
        "   | Categoria | Conteo | Porcentaje |\n"
        "7. Verifica que las categorías sean coherentes y pertinentes al contexto organizacional.\n"
    )



    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('47. Hábitos diarios')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_47_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_47_chart.png')

    return chart_url









# Obtener los datos de la pregunta_48
def generate_pregunta_48_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    habito_1 = PreguntaAbiertaHabitosMensuales.objects.values_list('habito_1', flat=True)
    habito_2 = PreguntaAbiertaHabitosMensuales.objects.values_list('habito_2', flat=True)
    respuestas = list(habito_1) + list(habito_2)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias que describan habitos mensuales, siempre 10, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              '"Mencione hábitos mensuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados."'
              'para crear las categorias usa los terminos en relacion con una jerga profesional en la PSICOLOGÍA ORGANIZACIONAL, haz un buen analisis del lenguaje'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'Los titulos de las categorias deben ser de más de una palabra para describir mejor el tema del que se habla'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'El formato debe ser estricto, los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , SIEMPRE se va a usar ese formato sin alteraciones, siempre. No agregues espacios adicionales o caracteres.'
              'no es necesario entregar más informacion ademas de las categorias, el conteo y el porcentaje en el formato solicitado.'
              'Por favor, que las categorias sean coherentes')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('48. Hábitos mensuales')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_48_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_48_chart.png')

    return chart_url















# Obtener los datos de la pregunta_49
def generate_pregunta_49_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    habito_1 = PreguntaAbiertaHabitosAnuales.objects.values_list('habito_1', flat=True)
    habito_2 = PreguntaAbiertaHabitosAnuales.objects.values_list('habito_2', flat=True)
    respuestas = list(habito_1) + list(habito_2)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias que describan habitos anuales, siempre 10, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              '"Mencione hábitos anuales que usted tenga dentro de la empresa, que le ayudan a obtener mejores resultados."'
              'para crear las categorias usa los terminos en relacion con una jerga profesional en la PSICOLOGÍA ORGANIZACIONAL'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'Los titulos de las categorias deben ser de más de una palabra para describir mejor el tema del que se habla'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'El formato debe ser estricto, los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , SIEMPRE se va a usar ese formato sin alteraciones, siempre.'
              'no es necesario entregar más informacion ademas de las categorias, el conteo y el porcentaje en el formato solicitado.'
              'Por favor, que las categorias sean coherentes')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    respuestaaaas = respuesta.replace('| |', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('49. Hábitos anuales')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_49_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_49_chart.png')

    return chart_url












# Obtener los datos de la pregunta_50
def generate_pregunta_50_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    defectos_persona_A = PreguntaAbiertaDefectosPersonas.objects.values_list('defectos_persona_A', flat=True)
    defectos_persona_B = PreguntaAbiertaDefectosPersonas.objects.values_list('defectos_persona_B', flat=True)
    defectos_persona_C = PreguntaAbiertaDefectosPersonas.objects.values_list('defectos_persona_C', flat=True)
    respuestas = list(defectos_persona_A) + list(defectos_persona_B) + list(defectos_persona_C)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias que describan sus defectos, siempre 10, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              '"Piense en personas que se destacan negativamente dentro de la empresa y señale sólo sus defectos."'
              'para crear las categorias usa los terminos en relacion con una jerga profesional en la PSICOLOGÍA ORGANIZACIONAL'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'Los titulos de las categorias deben ser de más de una palabra para describir mejor el tema del que se habla'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'El formato debe ser estricto, los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , SIEMPRE se va a usar ese formato sin alteraciones, siempre.'
              'no es necesario entregar más informacion ademas de las categorias, el conteo y el porcentaje en el formato solicitado.'
              'Por favor, que las categorias sean coherentes')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('50. Antihéroes')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_50_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_50_chart.png')

    return chart_url











# Obtener los datos de la pregunta_51
def generate_pregunta_51_chart(request):
    # Usamos el modelo generativo de la IA 
    modelo = genai.GenerativeModel('gemini-2.0-flash')
    # Configuramos la API KEY 
    GOOGLE_API_KEY='AIzaSyAlVNEfsQ1fC-YsGJ-BuWMhSBKkrbldUv4'
    genai.configure(api_key=GOOGLE_API_KEY)

    # Importamos las respuestas de la base de datos
    virtudes_persona_A = PreguntaAbiertaVirtudesPersonas.objects.values_list('virtudes_persona_A', flat=True)
    virtudes_persona_B = PreguntaAbiertaVirtudesPersonas.objects.values_list('virtudes_persona_B', flat=True)
    virtudes_persona_C = PreguntaAbiertaVirtudesPersonas.objects.values_list('virtudes_persona_C', flat=True)
    respuestas = list(virtudes_persona_A) + list(virtudes_persona_B) + list(virtudes_persona_C)

    # Convertimos las respuestas en un solo texto
    texto_respuestas = '|'.join(respuestas)

# Definimos el prompt
    prompt = ('genera 10 categorias que describan sus virtudes, siempre 10, basadas en los textos suministrados teniendo en cuenta que los titulos de las categorias deben estar relacionados a lo que es una organizacion y el contexto que es:'
              '"Piense personas que se destacan positivamente dentro de la empresa y señale sólo sus virtudes."'
              'para crear las categorias usa los terminos en relacion con una jerga profesional en la PSICOLOGÍA ORGANIZACIONAL'
              'usa todos los textos exceptuando los textos que solo dicen "Ninguna, .., N. A, ...,", por nada los vayas a usar para las categorias.'
              'Los titulos de las categorias deben ser de más de una palabra para describir mejor el tema del que se habla'
              'haz un conteo de todos los textos, ten en cuenta que estan separados por "|", dame un conteo y un porcentaje para cada categoria.'
              'El formato debe ser estricto, los datos siempre seran entregados en el siguiente formato, como si fuera una tabla: | Categoria | Conteo | Porcentaje | , SIEMPRE se va a usar ese formato sin alteraciones, siempre.'
              'no es necesario entregar más informacion ademas de las categorias, el conteo y el porcentaje en el formato solicitado.'
              'Por favor, que las categorias sean coherentes')


    # Generamos la respuesta basada en el prompt y el texto de las respuestas
    respuesta = modelo.generate_content(prompt + texto_respuestas)
    respuesta = respuesta.text

    # Rebajamos el tamaño de la respuesta de la IA 
    respuestaaaas = respuesta.replace('**', '|')
    respuestaaaas = respuesta.replace('*', '|')
    print (respuestaaaas)

    # Extraer las categorías, conteos y porcentajes
    categorias = re.findall(r'\| (.+?) \| (\d+) \| (\d+(?:[,.]\d+)?)% \|', respuestaaaas)
    print (categorias)

    # Convertir porcentajes a números
    conteos = [int(c[1]) for c in categorias]
    porcentajes = [float(c[2].replace(',', '.')) for c in categorias]

    # Convertir las categorías a una lista
    categorias = [c[0] for c in categorias]
 

    # Crear figura
    fig = Figure(figsize=(16, 6))
    ax = fig.add_subplot(111)

    # Crear gráfica de barras
    bars = ax.bar(categorias, conteos, color='lightgreen')

    # Añadir etiquetas y título
    ax.set_xlabel('Categoría')
    ax.set_ylabel('Conteo')
    ax.set_title('51. Héroes')
    ax.set_xticklabels(categorias, rotation=45, ha='right')

    # Ajustar márgenes
    fig.subplots_adjust(bottom=0.4, top=0.9) 

    # Mostrar porcentajes con decimales en las barras
    for bar, porcentaje in zip(bars, porcentajes):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05, f"{porcentaje:.2f}%", ha='center', fontsize=7)

    # Ajustar espacio entre etiquetas del eje x
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    # Guardar la imagen de la gráfica de barras en un directorio
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path_bar = os.path.join(chart_dir, 'pregunta_51_chart.png')
    fig.savefig(chart_path_bar, format='png')

    # Obtener la URL de la imagen de la gráfica de barras
    chart_url = os.path.join(settings.MEDIA_URL, 'charts', 'pregunta_51_chart.png')

    return chart_url








# Postulacion frente a premisas preguntas 52 a 61
# Función para obtener datos de las preguntas 52 a 61import os
def generar_tabla_liderazgo(request):
    """
    Obtiene los datos desde el modelo, los anonimiza, calcula frecuencias, 
    genera una tabla de resultados y agrupa los menos frecuentes en 'Otros'.
    Retorna la URL de la imagen generada con títulos de columna personalizados 
    y el resto del contenido en mayúsculas.
    """

    # Paso 1: Obtener los datos desde el modelo
    respuestas = PreguntaAbiertaCompaneros.objects.filter(
        fiesta_integracion__isnull=False, 
        defensa_intereses__isnull=False, 
        representante_directivas__isnull=False, 
        organizador_equipo_deportivo__isnull=False,
        organizador_equipos_trabajo__isnull=False,
        divulgacion_hechos__isnull=False,
        confianza_secreto__isnull=False,
        resolver_problemas__isnull=False,
        enseñanza_trabajo__isnull=False,
        lider_funcionario__isnull=False
    ).values(
        'fiesta_integracion', 'defensa_intereses', 'representante_directivas', 
        'organizador_equipo_deportivo', 'organizador_equipos_trabajo',
        'divulgacion_hechos', 'confianza_secreto', 'resolver_problemas', 
        'enseñanza_trabajo', 'lider_funcionario'
    )

    df = pd.DataFrame(respuestas)
    if df.empty:
        print("No se encontraron datos.")
        return None

    # Paso 2: Limpieza de datos
    import re
    from unidecode import unidecode

    palabras_excluidas = {'no', 'nose', 'nan', 'ninguno', 'ninguna', 'nada', 'null'}

    def limpiar_texto(texto):
        texto = str(texto).strip().lower()
        if texto in palabras_excluidas or texto == "":
            return ""
        texto = re.sub(r'[^\w\s]', '', texto)  # Eliminar puntuación
        texto = re.sub(r'\d+', '', texto)      # Eliminar números
        texto = unidecode(texto)               # Quitar acentos
        palabras = texto.split()
        texto_limpio = ' '.join([p for p in palabras if p.isalpha() and len(p) > 1])
        return texto_limpio

    for col in df.columns:
        df[col] = df[col].apply(limpiar_texto)

    def normalizar_texto(texto):
        return unidecode(texto).lower().strip() if texto else ""

    for col in df.columns:
        df[col] = df[col].apply(normalizar_texto)

    # Paso 3: Calcular frecuencias y porcentajes
    porcentajes_dict = {}
    for col in df.columns:
        df_col_filtrado = df[col][~df[col].isin(['', 'no', 'nose', 'nan', 'ninguno', 'ninguna', 'nada', 'null'])]
        frecuencias = df_col_filtrado.value_counts()
        total = frecuencias.sum()
        
        if total > 0:
            porcentajes = (frecuencias / total * 100).round(2)
            porcentajes_dict[col] = pd.DataFrame({
                'Nombre': frecuencias.index,
                'Porcentaje': porcentajes.values
            })
        else:
            porcentajes_dict[col] = pd.DataFrame({'Nombre': [], 'Porcentaje': []})

    # Paso 4: Construir la tabla sin agrupar
    tabla_final = pd.DataFrame()
    for col in porcentajes_dict:
        df_col = porcentajes_dict[col]
        df_col.columns = [f"{col}_Nombre", f"{col}_Porcentaje"]
        tabla_final = pd.concat([tabla_final, df_col], axis=1)

    # Función auxiliar para agrupar en "Otros" y mostrar solo los top_n valores
    def crear_tabla_agrupada(tabla, top_n=5):
        tabla_agrupada = pd.DataFrame()
        col_names = list(tabla.columns)
        for i in range(0, len(col_names), 2):
            nombre_col = col_names[i]
            porcentaje_col = col_names[i+1]
            df_temp = tabla[[nombre_col, porcentaje_col]].copy()
            df_temp.columns = ["Nombre", "Porcentaje"]

            df_temp = df_temp.sort_values(by="Porcentaje", ascending=False)
            df_top = df_temp.head(top_n)
            df_otros = df_temp.tail(len(df_temp) - top_n)
            porcentaje_otros = df_otros["Porcentaje"].sum()

            df_categoria = pd.DataFrame({
                f"{nombre_col}": list(df_top["Nombre"]) + ["Otros"],
                f"{porcentaje_col}": list(df_top["Porcentaje"]) + [porcentaje_otros]
            })

            if tabla_agrupada.empty:
                tabla_agrupada = df_categoria
            else:
                tabla_agrupada = pd.concat([tabla_agrupada.reset_index(drop=True),
                                            df_categoria.reset_index(drop=True)], axis=1)
        return tabla_agrupada

    # Paso 5: Agrupar con top_n=5
    tabla_agrupada = crear_tabla_agrupada(tabla_final, top_n=5)

    # Paso 5.1: Renombrar columnas para que aparezcan con los títulos solicitados
    # Cada par de columnas (Nombre/Porcentaje) se mapea a: "Social" y "%", "Profesional" y "%", etc.
    rename_dict = {
        'fiesta_integracion_Nombre': 'Social',
        'fiesta_integracion_Porcentaje': '%',
        'defensa_intereses_Nombre': 'Profesional',
        'defensa_intereses_Porcentaje': '%',
        'representante_directivas_Nombre': 'Gremial',
        'representante_directivas_Porcentaje': '%',
        'organizador_equipo_deportivo_Nombre': 'Deportivo',
        'organizador_equipo_deportivo_Porcentaje': '%',
        'organizador_equipos_trabajo_Nombre': 'Grupos de trabajo',
        'organizador_equipos_trabajo_Porcentaje': '%',
        'divulgacion_hechos_Nombre': 'Comunicador',
        'divulgacion_hechos_Porcentaje': '%',
        'confianza_secreto_Nombre': 'Confianza',
        'confianza_secreto_Porcentaje': '%',
        'resolver_problemas_Nombre': 'Conciliador',
        'resolver_problemas_Porcentaje': '%',
        'enseñanza_trabajo_Nombre': 'Tutor (Coach)',
        'enseñanza_trabajo_Porcentaje': '%',
        'lider_funcionario_Nombre': 'Carisma',
        'lider_funcionario_Porcentaje': '%'
    }
    tabla_agrupada.rename(columns=rename_dict, inplace=True)

    # Paso 6: Crear y guardar la figura con la nueva tabla
    fig, ax = plt.subplots(figsize=(20, 6))
    ax.axis('off')

    from pandas.plotting import table
    col_count = len(tabla_agrupada.columns)
    col_widths = [0.15] * col_count

    # Dibujamos la tabla base con los nuevos encabezados
    tabla_plot = table(ax, tabla_agrupada, loc='center', cellLoc='center', colWidths=col_widths)
    tabla_plot.auto_set_font_size(False)
    tabla_plot.set_fontsize(12)
    tabla_plot.scale(1.2, 1.2)

    # Paso 7: Añadir colores y estilos
    header_color = '#4CAF50'
    category_colors = [
        '#D9EAD3', '#FCE5CD', '#CFE2F3', '#FFF2CC', '#F4CCCC',
        '#D2E4BD', '#F9CB9C', '#C9DAF8', '#FFE599', '#EA9999',
    ]

    # Paso 8: Dar formato a celdas (mantener encabezados tal como los definimos y poner filas en mayúsculas)
    for (row, col), cell in tabla_plot.get_celld().items():
        if row == 0:
            # Encabezado: color verde y texto blanco en negrita
            cell.set_facecolor(header_color)
            cell.set_text_props(color='white', weight='bold')
        else:
            # Resto de filas: texto en mayúsculas y colores alternados por categoría
            current_text = cell.get_text().get_text()
            cell.get_text().set_text(current_text.upper())  # mayúsculas para los datos
            color_index = (col // 2) % len(category_colors)
            cell.set_facecolor(category_colors[color_index])
            cell.set_text_props(color='black')

    # Guardar la imagen
    chart_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(chart_dir, exist_ok=True)
    chart_path = os.path.join(chart_dir, 'tabla_liderazgo.png')
    fig.savefig(chart_path, bbox_inches='tight', dpi=300)

    # Retornar la URL de la imagen generada
    table_url = os.path.join(settings.MEDIA_URL, 'charts', 'tabla_liderazgo.png')
    return table_url









#Sintesis
#Se enlaaza en las funciones
def sintesis_administracion(request):
    print("Inicio de sintesis_administracion")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_11_chart,
        generate_pregunta_13_chart,
        generate_pregunta_14_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])  # Ajustar anchos
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_admin.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_admin.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_admin.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_admin.png')

    print("Fin de sintesis_administracion")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }
    







def sintesis_calidad(request):
    print("Inicio de sintesis_calidad")
    funciones = [
        generate_pregunta_11_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_16_chart,
        generate_pregunta_31_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            # Se asume que los índices 2 y 5 corresponden a la suma de acuerdos y desacuerdos respectivamente
            porcentaje_acuerdos = percentages[2]
            porcentaje_desacuerdos = percentages[5]
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de los porcentajes más altos
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_calidad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_calidad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_calidad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_calidad.png')

    print("Fin de sintesis_calidad")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }
    








def sintesis_cambio(request):
    print("Inicio de sintesis_cambio")
    funciones = [
        generate_pregunta_7_chart,
        generate_pregunta_13_chart,
        generate_pregunta_16_chart,
        generate_pregunta_22_chart,
        generate_pregunta_25_chart,
        generate_pregunta_27_chart,
        generate_pregunta_29_chart,
        generate_pregunta_32_chart,
        generate_pregunta_33_chart,
        generate_pregunta_34_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de los porcentajes más altos
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_cambio.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_cambio.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_cambio.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_cambio.png')

    print("Fin de sintesis_cambio")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }





def sintesis_comunicacion(request):
    print("Inicio de sintesis_comunicacion")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_10_chart,
        generate_pregunta_13_chart,
        generate_pregunta_15_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_19_chart,
        generate_pregunta_26_chart,
        generate_pregunta_27_chart,
        generate_pregunta_35_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_comunicacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_comunicacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_comunicacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_comunicacion.png')

    print("Fin de sintesis_comunicacion")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }







def sintesis_conflicto_negociacion(request):
    print("Inicio de sintesis_conflicto_negociacion")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_13_chart,
        generate_pregunta_15_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_19_chart,
        generate_pregunta_21_chart,
        generate_pregunta_26_chart,
        generate_pregunta_27_chart,
        generate_pregunta_35_chart,
        generate_pregunta_38_chart,
        generate_pregunta_39_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', 
                           colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_conflicto_negociacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_conflicto_negociacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', 
                    label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_conflicto_negociacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_conflicto_negociacion.png')

    print("Fin de sintesis_conflicto_negociacion")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }





def sintesis_delegacion_confianza_functional(request):
    print("Inicio de sintesis_delegacion_confianza_functional")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_11_chart,
        generate_pregunta_13_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_23_chart,
        generate_pregunta_27_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_delegacion_confianza.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_delegacion_confianza.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_delegacion_confianza.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_delegacion_confianza.png')

    print("Fin de sintesis_delegacion_confianza_functional")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }







def sintesis_Eficacia_Coherencia(request):
    print("Inicio de sintesis_Eficacia_Coherencia")
    funciones = [
        generate_pregunta_7_chart,
        generate_pregunta_11_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_17_chart,
        generate_pregunta_22_chart,
        generate_pregunta_27_chart,
        generate_pregunta_31_chart,
        generate_pregunta_34_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_eficacia_coherencia.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_eficacia_coherencia.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos (Eficacia/Coherencia)")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_eficacia_coherencia.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_eficacia_coherencia.png')

    print("Fin de sintesis_Eficacia_Coherencia")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_eficiencia(request):
    """
    Sintesis_Eficiencia: Ahorro Vs Desperdicio interno
    """
    print("Inicio de sintesis_Eficiencia: Ahorro Vs Desperdicio interno")
    funciones = [
        generate_pregunta_11_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_14_chart,
        generate_pregunta_17_chart,
        generate_pregunta_22_chart,
        generate_pregunta_33_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_eficiencia.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_eficiencia.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Ahorro Vs Desperdicio Interno")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_eficiencia.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_eficiencia.png')

    print("Fin de sintesis_Eficiencia: Ahorro Vs Desperdicio interno")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }







def sintesis_Liderazgo(request):
    print("Inicio de sintesis_Liderazgo")
    funciones = [
        generate_pregunta_5_chart,
        generate_pregunta_7_chart,
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_10_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_15_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_19_chart,
        generate_pregunta_27_chart,
        generate_pregunta_29_chart,
        generate_pregunta_31_chart,
        generate_pregunta_35_chart,
        generate_pregunta_37_chart,
        generate_pregunta_38_chart,
        generate_pregunta_39_chart,
        generate_pregunta_50_chart,
        generate_pregunta_51_chart,
        #generate_pregunta_62_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', 
                           colWidths=[0.15, 0.20, 0.20, 0.45])  # Ajustar anchos
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_liderazgo.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_liderazgo.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_liderazgo.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_liderazgo.png')

    print("Fin de sintesis_Liderazgo")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_Orientación(request):
    print("Inicio de sintesis_Orientación")
    funciones = [
        generate_pregunta_7_chart,
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_11_chart,
        generate_pregunta_13_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_22_chart,
        generate_pregunta_24_chart,
        generate_pregunta_31_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_orientacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_orientacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_orientacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_orientacion.png')

    print("Fin de sintesis_Orientación")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_poder_autoridad(request):
    print("Inicio de sintesis_Poder/Autoridad")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_15_chart,
        generate_pregunta_17_chart,
        generate_pregunta_26_chart,
        generate_pregunta_27_chart,
        generate_pregunta_37_chart,
        generate_pregunta_38_chart,
        generate_pregunta_39_chart,
        generate_pregunta_44_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']]
          for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', 
                           colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_poder_autoridad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_poder_autoridad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--',
                    label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_poder_autoridad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_poder_autoridad.png')

    print("Fin de sintesis_Poder/Autoridad")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_trabajo_equipo(request):
    print("Inicio de sintesis_trabajo_equipo")
    funciones = [
        generate_pregunta_5_chart,
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_10_chart,
        generate_pregunta_11_chart,
        generate_pregunta_12_chart,
        generate_pregunta_13_chart,
        generate_pregunta_15_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_19_chart,
        generate_pregunta_20_chart,
        generate_pregunta_21_chart,
        generate_pregunta_26_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_33_chart,
        generate_pregunta_35_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_trabajo_equipo.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_trabajo_equipo.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_trabajo_equipo.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_trabajo_equipo.png')

    print("Fin de sintesis_trabajo_equipo")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_Calidez_Frialdad(request):
    print("Inicio de sintesis_Calidez_Frialdad")
    funciones = [
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_10_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_19_chart,
        generate_pregunta_20_chart,
        generate_pregunta_21_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_calidez_frialdad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_calidez_frialdad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_calidez_frialdad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_calidez_frialdad.png')

    print("Fin de sintesis_Calidez_Frialdad")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_Autonomia_Control(request):
    print("Inicio de sintesis_Autonomia_Control")
    funciones = [
        generate_pregunta_1_chart,
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_11_chart,
        generate_pregunta_17_chart,
        generate_pregunta_22_chart,
        generate_pregunta_23_chart,
        generate_pregunta_25_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_autonomia_control.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_autonomia_control.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_autonomia_control.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_autonomia_control.png')

    print("Fin de sintesis_Autonomia_Control")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_Desarrollo_de_Carrera(request):
    print("Inicio de sintesis_Desarrollo de Carrera")
    funciones = [
        generate_pregunta_7_chart,
        generate_pregunta_16_chart,
        generate_pregunta_22_chart,
        generate_pregunta_24_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_41_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(
        cellText=tabla_data,
        cellLoc='left',
        loc='center',
        colWidths=[0.15, 0.20, 0.20, 0.45]  # Ajustar anchos
    )
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_desarrollo.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_desarrollo.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_desarrollo.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_desarrollo.png')

    print("Fin de sintesis_Desarrollo de Carrera")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_Iniciativa(request):
    print("Inicio de sintesis_Iniciativa")
    funciones = [
        generate_pregunta_2_chart,
        generate_pregunta_3_chart,
        generate_pregunta_23_chart,
        generate_pregunta_25_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_iniciativa.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_iniciativa.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_iniciativa.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_iniciativa.png')

    print("Fin de sintesis_Iniciativa")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_Innovación(request):
    print("Inicio de sintesis_Innovación")
    funciones = [
        generate_pregunta_11_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_22_chart,
        generate_pregunta_23_chart,
        generate_pregunta_25_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_innovacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_innovacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_innovacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_innovacion.png')

    print("Fin de sintesis_Innovación")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }










def sintesis_Libertad_de_pensamiento(request):
    print("Inicio de sintesis_Libertad_de_pensamiento")
    funciones = [
        generate_pregunta_1_chart,
        generate_pregunta_15_chart,
        generate_pregunta_17_chart,
        generate_pregunta_19_chart,
        generate_pregunta_23_chart,
        generate_pregunta_24_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_libertad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_libertad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_libertad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_libertad.png')

    print("Fin de sintesis_Libertad_de_pensamiento")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_Motivación(request):
    print("Inicio de sintesis_Motivación")
    funciones = [
        generate_pregunta_7_chart,
        generate_pregunta_8_chart,
        generate_pregunta_9_chart,
        generate_pregunta_10_chart,
        generate_pregunta_15_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_29_chart,
        generate_pregunta_36_chart,
        generate_pregunta_37_chart,
        generate_pregunta_38_chart,
        generate_pregunta_39_chart,
        generate_pregunta_41_chart,
        generate_pregunta_43_chart,
        #generate_pregunta_52_chart,  # Rango 52-61
        #generate_pregunta_53_chart,
        #generate_pregunta_54_chart,
        #generate_pregunta_55_chart,
        #generate_pregunta_56_chart,
        #generate_pregunta_57_chart,
        #generate_pregunta_58_chart,
        #generate_pregunta_59_chart,
        #generate_pregunta_60_chart,
        #generate_pregunta_61_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_motivacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_motivacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_motivacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_motivacion.png')

    print("Fin de sintesis_Motivación")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_Pertenencia(request):
    print("Inicio de sintesis_Pertenencia")
    ffunciones = [
        generate_pregunta_12_chart,
        generate_pregunta_15_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_24_chart,
        generate_pregunta_25_chart,
        generate_pregunta_26_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_29_chart,
        generate_pregunta_30_chart,
        generate_pregunta_31_chart,
        generate_pregunta_32_chart,
        generate_pregunta_33_chart,
        generate_pregunta_34_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in ffunciones])
    datos = []
    highest_percentages = []

    for func in ffunciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_pertenencia.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_pertenencia.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_pertenencia.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_pertenencia.png')

    print("Fin de sintesis_Pertenencia")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }








def sintesis_satisfaccion_insatisfaccion_con_la_organizacion(request):
    print("Inicio de sintesis_satisfaccion_insatisfaccion_con_la_organizacion")
    funciones = [
        generate_pregunta_11_chart,
        generate_pregunta_12_chart,
        generate_pregunta_15_chart,
        generate_pregunta_16_chart,
        generate_pregunta_17_chart,
        generate_pregunta_18_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_29_chart,
        generate_pregunta_32_chart,
        generate_pregunta_33_chart,
        generate_pregunta_34_chart,
        generate_pregunta_37_chart,
        generate_pregunta_38_chart,
        generate_pregunta_39_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_satisfaccion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_satisfaccion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    
    line_path = os.path.join(tabla_dir, 'sintesis_linea_satisfaccion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_satisfaccion.png')

    print("Fin de sintesis_satisfaccion_insatisfaccion_con_la_organizacion")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_La_organizacion_aprovecha_las_competencias_individuales(request):
    print("Inicio de sintesis_La_organizacion_aprovecha_las_competencias_individuales")
    funciones = [
        generate_pregunta_2_chart,
        generate_pregunta_22_chart,
        generate_pregunta_23_chart,
        generate_pregunta_24_chart,
        generate_pregunta_25_chart,
        generate_pregunta_34_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20, 0.45])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_organizacion.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_organizacion.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_organizacion.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_organizacion.png')

    print("Fin de sintesis_La_organizacion_aprovecha_las_competencias_individuales")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }









def sintesis_Equidad(request):
    print("Inicio de sintesis_Equidad en la carga laboral")
    funciones = [
        generate_pregunta_17_chart,
        generate_pregunta_27_chart,
        generate_pregunta_28_chart,
        generate_pregunta_32_chart,
        generate_pregunta_33_chart,
        generate_pregunta_36_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', 
                           colWidths=[0.15, 0.20, 0.20, 0.45])  # Ajustar anchos
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_equidad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_equidad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', 
                    label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    
    line_path = os.path.join(tabla_dir, 'sintesis_linea_equidad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_equidad.png')

    print("Fin de sintesis_Equidad en la carga laboral")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }










def sintesis_Percepción_de_la_normatividad(request):
    print("Inicio de sintesis_Percepción_de_la_normatividad")
    funciones = [
        generate_pregunta_1_chart,
        generate_pregunta_2_chart,
        generate_pregunta_4_chart,
        generate_pregunta_5_chart,
        generate_pregunta_15_chart,
        generate_pregunta_17_chart,
        generate_pregunta_23_chart,
        generate_pregunta_25_chart,
        generate_pregunta_26_chart,
        generate_pregunta_40_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(
        cellText=tabla_data,
        cellLoc='left',
        loc='center',
        colWidths=[0.15, 0.20, 0.20, 0.45]
    )  # Ajustar anchos
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_normatividad.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_normatividad.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("24. - Normas según la flexibilidad percibida: Flexible - Rígida ")
    ax_line.set_xlabel("Items")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_normatividad.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_normatividad.png')

    print("Fin de sintesis_Percepción_de_la_normatividad")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
        

    }











#def sintesis_Abiertas(request):
    print("Inicio de sintesis_Abiertas")
    funciones = [
        generate_pregunta_42_chart,
        generate_pregunta_45_chart,
        generate_pregunta_46_chart,
        generate_pregunta_47_chart,
        generate_pregunta_48_chart,
        generate_pregunta_49_chart
    ]
    print("Funciones definidas:", [func.__name__ for func in funciones])
    datos = []
    highest_percentages = []

    for func in funciones:
        try:
            print(f"Ejecutando {func.__name__}")
            result = func(request)
            print(f"Resultado de {func.__name__}: {result}")
            percentages = result['percentages_ordenadas']
            porcentaje_acuerdos = percentages[2]  # Suma de acuerdos
            porcentaje_desacuerdos = percentages[5]  # Suma de desacuerdos
            inferencia = result['inferencia']
            highest_percentage = max(porcentaje_acuerdos, porcentaje_desacuerdos)
            datos.append({
                'pregunta': func.__name__.replace('generate_pregunta_', '').replace('_chart', ''),
                'highest_percentage': highest_percentage,
                'inferencia': inferencia
            })
            highest_percentages.append(highest_percentage)
        except Exception as e:
            print(f"Error en {func.__name__}: {e}")

    # Calcular el promedio de highest_percentages
    promedio_tendencias = np.mean(highest_percentages) if highest_percentages else 0
    print(f"Promedio de highest_percentages: {promedio_tendencias:.1f}%")

    # Generar la tabla con la nueva columna
    fig_tabla = Figure(figsize=(16, 6))  # Aumentar el ancho para la nueva columna
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Nivel de expresión del ítem", "Promedio de tendencias", "Tendencias/Indicadores de cumplimiento"],
        *[[f"Pregunta {d['pregunta']}", f"{d['highest_percentage']:.1f}%", f"{promedio_tendencias:.1f}%", d['inferencia']] for d in datos]
    ]
    tabla = ax_tabla.table(
        cellText=tabla_data, cellLoc='left', loc='center',
        colWidths=[0.15, 0.20, 0.20, 0.45]
    )  # Ajustar anchos
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_abiertas.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f"Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"Error al guardar la tabla: {e}")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_abiertas.png')

    # Generar la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [f"Pregunta {d['pregunta']}" for d in datos]
    ax_line.plot(preguntas, highest_percentages, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_tendencias, color='red', linestyle='--', label=f'Promedio: {promedio_tendencias:.1f}%')
    ax_line.set_title("Tendencias de Acuerdos y Desacuerdos")
    ax_line.set_xlabel("Preguntas")
    ax_line.set_ylabel("Porcentaje más alto")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_abiertas.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f"Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_abiertas.png')

    print("Fin de sintesis_Abiertas")
    return {
        'tabla_url': tabla_url,
        'linea_url': linea_url,
        'promedio_tendencias': promedio_tendencias,
    }








#SINTESIS DE LOS 24 INDICADORES 
def sintesis_general(request):
    # Mensaje inicial para confirmar que se llamó a la función
    print(">>> Inicio de sintesis_general")
    
    # Lista de funciones de síntesis
    funciones = [
        sintesis_administracion,
        sintesis_calidad,
        sintesis_cambio,
        sintesis_comunicacion,
        sintesis_conflicto_negociacion,
        sintesis_delegacion_confianza_functional,
        sintesis_Eficacia_Coherencia,
        sintesis_eficiencia,
        sintesis_Liderazgo,
        sintesis_Orientación,
        sintesis_poder_autoridad,
        sintesis_trabajo_equipo,
        sintesis_Calidez_Frialdad,
        sintesis_Autonomia_Control,
        sintesis_Desarrollo_de_Carrera,
        sintesis_Iniciativa,
        sintesis_Innovación,
        sintesis_Libertad_de_pensamiento,
        sintesis_Motivación,
        sintesis_Pertenencia,
        sintesis_satisfaccion_insatisfaccion_con_la_organizacion,
        sintesis_La_organizacion_aprovecha_las_competencias_individuales,
        sintesis_Equidad,
        sintesis_Percepción_de_la_normatividad
    ]
    
    # Lista de títulos asignados explícitamente según el orden deseado
    titulos = [
        "Administración",
        "Calidad",
        "Cambio",
        "Comunicación",
        "Conflicto/Negociación",
        "Delegación: Confianza funcional",
        "Eficacia: Coherencia",
        "Eficiencia: Ahorro Vs. Desperdicio interno",
        "Liderazgo",
        "Orientación",
        "Poder/Autoridad",
        "Trabajo en Equipo",
        "Calidez/Frialdad",
        "Autonomía/Autocontrol",
        "Desarrollo de Carrera",
        "Iniciativa",
        "Innovación",
        "Libertad de pensamiento",
        "Motivación",
        "Pertenencia",
        "Satisfacción/Insatisfacción con la organización",
        "Competencias individuales",
        "Equidad en la carga laboral",
        "Flexibilidad/Rigidez"
    ]
    
    datos = []
    indices_aceptacion = []
    
    # Procesamiento de cada función
    for i, func in enumerate(funciones):
        print(f">>> Procesando función: {func.__name__}")
        try:
            result = func(request)
            print(f">>> Resultado de {func.__name__}: {result}")
            if 'promedio_tendencias' not in result:
                print(f"!!! Advertencia: 'promedio_tendencias' no se encuentra en el resultado de {func.__name__}")
            else:
                indice = result['promedio_tendencias']
                datos.append({
                    'pregunta': titulos[i],  # Se asigna el título correspondiente del arreglo 'titulos'
                    'indice_aceptacion': indice,
                })
                indices_aceptacion.append(indice)
        except Exception as e:
            print(f"!!! Error en {func.__name__}: {e}")
    
    if not datos:
        print("!!! No se obtuvieron datos de ninguna función.")
    else:
        print(">>> Datos recopilados:", datos)
    
    # Cálculo del promedio para depuración
    if indices_aceptacion:
        promedio_indice = np.mean(indices_aceptacion)
    else:
        promedio_indice = 0
    print(f">>> Promedio del Índice de aceptación (para depuración): {promedio_indice:.1f}%")
    
    # Directorio para guardar imágenes
    tabla_dir = os.path.join(settings.BASE_DIR, 'media', 'charts')
    os.makedirs(tabla_dir, exist_ok=True)
    print(f">>> Directorio para guardar imágenes: {tabla_dir}")
    
    # Generación de la tabla sin la columna de Tendencias/Indicadores
    fig_tabla = Figure(figsize=(16, 6))
    ax_tabla = fig_tabla.add_subplot(111)
    ax_tabla.axis('off')
    tabla_data = [
        ["Ítem", "Índice de aceptación", "Promedio (depuración)"],
        *[[d['pregunta'], f"{d['indice_aceptacion']:.1f}%", f"{promedio_indice:.1f}%"] for d in datos]
    ]
    tabla = ax_tabla.table(cellText=tabla_data, cellLoc='left', loc='center', colWidths=[0.15, 0.20, 0.20])
    tabla.auto_set_font_size(False)
    tabla.set_fontsize(10)
    tabla.scale(1, 1.5)
    tabla_path = os.path.join(tabla_dir, 'sintesis_tabla_general.png')
    try:
        fig_tabla.savefig(tabla_path, bbox_inches='tight')
        print(f">>> Tabla guardada en: {tabla_path}")
    except Exception as e:
        print(f"!!! Error al guardar la tabla: {e}")
    
    # Verificar la existencia del archivo de la tabla
    if os.path.exists(tabla_path):
        print(">>> Verificación: La tabla se ha creado correctamente.")
    else:
        print("!!! Verificación: La tabla NO se ha creado.")
    tabla_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_tabla_general.png')
    
    # Generación de la gráfica de líneas
    fig_line = Figure(figsize=(16, 6))
    ax_line = fig_line.add_subplot(111)
    preguntas = [d['pregunta'] for d in datos]
    ax_line.plot(preguntas, indices_aceptacion, marker='o', linestyle='-', color='blue')
    ax_line.axhline(y=promedio_indice, color='red', linestyle='--', label=f'Promedio: {promedio_indice:.1f}%')
    ax_line.set_title("Índice de aceptación por dimensión")
    ax_line.set_xlabel("Dimensiones")
    ax_line.set_ylabel("Índice de aceptación (%)")
    ax_line.legend()
    ax_line.set_xticks(range(len(preguntas)))
    ax_line.set_xticklabels(preguntas, rotation=45, ha='right')
    line_path = os.path.join(tabla_dir, 'sintesis_linea_general.png')
    try:
        fig_line.savefig(line_path, bbox_inches='tight')
        print(f">>> Gráfica guardada en: {line_path}")
    except Exception as e:
        print(f"!!! Error al guardar la gráfica: {e}")
    plt.close(fig_line)
    
    # Verificar la existencia del archivo de la gráfica
    if os.path.exists(line_path):
        print(">>> Verificación: La gráfica se ha creado correctamente.")
    else:
        print("!!! Verificación: La gráfica NO se ha creado.")
    linea_url = os.path.join(settings.MEDIA_URL, 'charts', 'sintesis_linea_general.png')
    
    print(">>> Fin de sintesis_general")
    return {'tabla_url': tabla_url, 'linea_url': linea_url}












def view_results(request):

    # Llamar a la función para generar la gráfica de barras y obtener la URL de la imagen
    age_chart_url = generate_age_bar_chart(request)
    gender_chart_url = generate_gender_bar_chart(request)
    area_empresa_chart_url = generate_area_empresa_chart(request)
    antiguedad_empresa_chart_url = generate_antiguedad_empresa_chart(request)
    table_url = tabla_datos(request)
    pregunta_1_data = generate_pregunta_1_chart(request)
    pregunta_1_chart_url = pregunta_1_data['chart_url']
    
    pregunta_2_data = generate_pregunta_2_chart(request)
    pregunta_2_chart_url = pregunta_2_data['chart_url']
    
    pregunta_3_data = generate_pregunta_3_chart(request)
    pregunta_3_chart_url = pregunta_3_data['chart_url']
    
    pregunta_4_data = generate_pregunta_4_chart(request)
    pregunta_4_chart_url = pregunta_4_data['chart_url']
    
    pregunta_5_data = generate_pregunta_5_chart(request)
    pregunta_5_chart_url = pregunta_5_data['chart_url']
    
    pregunta_6_data = generate_pregunta_6_chart(request)
    pregunta_6_chart_url = pregunta_6_data['chart_url']
    
    pregunta_7_data = generate_pregunta_7_chart(request)
    pregunta_7_chart_url = pregunta_7_data['chart_url']
    
    pregunta_8_data = generate_pregunta_8_chart(request)
    pregunta_8_chart_url = pregunta_8_data['chart_url']
    
    pregunta_9_data = generate_pregunta_9_chart(request)
    pregunta_9_chart_url = pregunta_9_data['chart_url']
    
    pregunta_10_data = generate_pregunta_10_chart(request)
    pregunta_10_chart_url = pregunta_10_data['chart_url']
    
    pregunta_11_data = generate_pregunta_11_chart(request)
    pregunta_11_chart_url = pregunta_11_data['chart_url']
    
    pregunta_12_data = generate_pregunta_12_chart(request)
    pregunta_12_chart_url = pregunta_12_data['chart_url']
    
    pregunta_13_data = generate_pregunta_13_chart(request)
    pregunta_13_chart_url = pregunta_13_data['chart_url']
    
    pregunta_14_data = generate_pregunta_14_chart(request)
    pregunta_14_chart_url = pregunta_14_data['chart_url']
    
    pregunta_15_data = generate_pregunta_15_chart(request)
    pregunta_15_chart_url = pregunta_15_data['chart_url']
    
    pregunta_16_data = generate_pregunta_16_chart(request)
    pregunta_16_chart_url = pregunta_16_data['chart_url']
    
    pregunta_17_data = generate_pregunta_17_chart(request)
    pregunta_17_chart_url = pregunta_17_data['chart_url']
    
    pregunta_18_data = generate_pregunta_18_chart(request)
    pregunta_18_chart_url = pregunta_18_data['chart_url']
    
    pregunta_19_data = generate_pregunta_19_chart(request)
    pregunta_19_chart_url = pregunta_19_data['chart_url']
    
    pregunta_20_data = generate_pregunta_20_chart(request)
    pregunta_20_chart_url = pregunta_20_data['chart_url']
    
    pregunta_21_data = generate_pregunta_21_chart(request)
    pregunta_21_chart_url = pregunta_21_data['chart_url']
    
    pregunta_22_data = generate_pregunta_22_chart(request)
    pregunta_22_chart_url = pregunta_22_data['chart_url']
    
    pregunta_23_data = generate_pregunta_23_chart(request)
    pregunta_23_chart_url = pregunta_23_data['chart_url']
    
    pregunta_24_data = generate_pregunta_24_chart(request)
    pregunta_24_chart_url = pregunta_24_data['chart_url']
    
    pregunta_25_data = generate_pregunta_25_chart(request)
    pregunta_25_chart_url = pregunta_25_data['chart_url']
    
    pregunta_26_data = generate_pregunta_26_chart(request)
    pregunta_26_chart_url = pregunta_26_data['chart_url']
    
    pregunta_27_data = generate_pregunta_27_chart(request)
    pregunta_27_chart_url = pregunta_27_data['chart_url']
    
    pregunta_28_data = generate_pregunta_28_chart(request)
    pregunta_28_chart_url = pregunta_28_data['chart_url']
    
    pregunta_29_data = generate_pregunta_29_chart(request)
    pregunta_29_chart_url = pregunta_29_data['chart_url']
    
    pregunta_30_data = generate_pregunta_30_chart(request)
    pregunta_30_chart_url = pregunta_30_data['chart_url']
    
    pregunta_31_data = generate_pregunta_31_chart(request)
    pregunta_31_chart_url = pregunta_31_data['chart_url']
    
    pregunta_32_data = generate_pregunta_32_chart(request)
    pregunta_32_chart_url = pregunta_32_data['chart_url']
    
    pregunta_33_data = generate_pregunta_33_chart(request)
    pregunta_33_chart_url = pregunta_33_data['chart_url']
    
    pregunta_34_data = generate_pregunta_34_chart(request)
    pregunta_34_chart_url = pregunta_34_data['chart_url']
    
    pregunta_35_data = generate_pregunta_35_chart(request)
    pregunta_35_chart_url = pregunta_35_data['chart_url']
    
    pregunta_36_data = generate_pregunta_36_chart(request)
    pregunta_36_chart_url = pregunta_36_data['chart_url']
    
    pregunta_37_data = generate_pregunta_37_chart(request)
    pregunta_37_chart_url = pregunta_37_data['chart_url']
    
    pregunta_38_data = generate_pregunta_38_chart(request)
    pregunta_38_chart_url = pregunta_38_data['chart_url']
    
    pregunta_39_data = generate_pregunta_39_chart(request)
    pregunta_39_chart_url = pregunta_39_data['chart_url']
    
    pregunta_40_data = generate_pregunta_40_chart(request)
    pregunta_40_chart_url = pregunta_40_data['chart_url']
    
    pregunta_41_data = generate_pregunta_41_chart(request)
    pregunta_41_chart_url = pregunta_41_data['chart_url']

    pregunta_42_chart_bar3_url = procesar_respuestas3(request)
    pregunta_43_opcion_1_chart_url = generate_pregunta_43_opcion_1_chart(request)
    pregunta_43_opcion_2_chart_url = generate_pregunta_43_opcion_2_chart(request)
    pregunta_43_opcion_3_chart_url = generate_pregunta_43_opcion_3_chart(request)
    pregunta_43_chart_url = generate_pregunta_43_chart(request)
    pregunta_44_opcion_1_chart_url = generate_pregunta_44_opcion_1_chart(request)
    pregunta_44_opcion_2_chart_url = generate_pregunta_44_opcion_2_chart(request)
    pregunta_44_opcion_3_chart_url = generate_pregunta_44_opcion_3_chart(request)
    pregunta_44_opcion_4_chart_url = generate_pregunta_44_opcion_4_chart(request)
    pregunta_44_chart_url = generate_pregunta_44_chart(request)
    pregunta_45_chart_url = generate_pregunta_45_chart(request)
    pregunta_46_chart_url = generate_pregunta_46_chart(request)
    pregunta_47_chart_url = generate_pregunta_47_chart(request)
    pregunta_48_chart_url = generate_pregunta_48_chart(request)
    pregunta_49_chart_url = generate_pregunta_49_chart(request)
    pregunta_50_chart_url = generate_pregunta_50_chart(request)
    pregunta_51_chart_url = generate_pregunta_51_chart(request)
    preguntas_52_a_61_table_url = generar_tabla_liderazgo(request)
    # Sionetesis de administracion
    sintesis_data = sintesis_administracion(request)
    tabla_admin_url = sintesis_data['tabla_url']
    linea_admin_url = sintesis_data['linea_url']
    # Síntesis de calidad 
    sintesis_calidad_data = sintesis_calidad(request)
    tabla_calidad_url = sintesis_calidad_data['tabla_url']
    linea_calidad_url = sintesis_calidad_data['linea_url']
    # Sintesis de cambio 
    sintesis_cambio_data = sintesis_cambio(request)
    tabla_cambio_url = sintesis_cambio_data['tabla_url']
    linea_cambio_url = sintesis_cambio_data['linea_url']
    # Síntesis de comunicación
    sintesis_comunicacion_data = sintesis_comunicacion(request)
    tabla_comunicacion_url = sintesis_comunicacion_data['tabla_url']
    linea_comunicacion_url = sintesis_comunicacion_data['linea_url']
    # Sintesis de conflicto y negociación
    sintesis_conflicto_negociacion_data = sintesis_conflicto_negociacion(request)
    tabla_conflicto_negociacion_url = sintesis_conflicto_negociacion_data['tabla_url']
    linea_conflicto_negociacion_url = sintesis_conflicto_negociacion_data['linea_url']
    # Sintesis de delegación y confianza 
    sintesis_delegacion_confianza_data = sintesis_delegacion_confianza_functional(request)
    tabla_delegacion_confianza_url = sintesis_delegacion_confianza_data['tabla_url']
    linea_delegacion_confianza_url = sintesis_delegacion_confianza_data['linea_url']
    # Sintesis de Eficacia/Coherencia
    sintesis_eficacia_coherencia_data = sintesis_Eficacia_Coherencia(request)
    tabla_eficacia_coherencia_url = sintesis_eficacia_coherencia_data['tabla_url']
    linea_eficacia_coherencia_url = sintesis_eficacia_coherencia_data['linea_url']
    # Sintesis de Eficiencia: Ahorro Vs Desperdicio interno
    sintesis_eficiencia_data = sintesis_eficiencia(request)
    tabla_eficiencia_url = sintesis_eficiencia_data['tabla_url']
    linea_eficiencia_url = sintesis_eficiencia_data['linea_url']
    # Síntesis de Liderazgo
    sintesis_liderazgo_data = sintesis_Liderazgo(request)
    tabla_liderazgo_url = sintesis_liderazgo_data['tabla_url']
    linea_liderazgo_url = sintesis_liderazgo_data['linea_url']
    # Síntesis de Orientación
    sintesis_orientacion_data = sintesis_Orientación(request)
    sintesis_orientacion_tabla_url = sintesis_orientacion_data['tabla_url']
    sintesis_orientacion_linea_url = sintesis_orientacion_data['linea_url']
    # Síntesis de Poder/Autoridad
    sintesis_poder_autoridad_data = sintesis_poder_autoridad(request)
    tabla_poder_autoridad_url = sintesis_poder_autoridad_data['tabla_url']
    linea_poder_autoridad_url = sintesis_poder_autoridad_data['linea_url']
    # Síntesis de Trabajo en Equipo
    sintesis_trabajo_equipo_data = sintesis_trabajo_equipo(request)
    tabla_trabajo_equipo_url = sintesis_trabajo_equipo_data['tabla_url']
    linea_trabajo_equipo_url = sintesis_trabajo_equipo_data['linea_url']
    # Síntesis de Calidez/Frialdad
    sintesis_calidez_frialdad_data = sintesis_Calidez_Frialdad(request)
    tabla_calidez_frialdad_url = sintesis_calidez_frialdad_data['tabla_url']
    linea_calidez_frialdad_url = sintesis_calidez_frialdad_data['linea_url']
    # Síntesis de Autonomía/Control
    sintesis_autonomia_control_data = sintesis_Autonomia_Control(request)
    tabla_autonomia_control_url = sintesis_autonomia_control_data['tabla_url']
    linea_autonomia_control_url = sintesis_autonomia_control_data['linea_url']
    # Síntesis de Desarrollo de Carrera
    sintesis_desarrollo_data = sintesis_Desarrollo_de_Carrera(request)
    tabla_desarrollo_url = sintesis_desarrollo_data['tabla_url']
    linea_desarrollo_url = sintesis_desarrollo_data['linea_url']
    # Síntesis de iniciativa
    sintesis_iniciativa_data = sintesis_Iniciativa(request)
    tabla_iniciativa_url = sintesis_iniciativa_data['tabla_url']
    linea_iniciativa_url = sintesis_iniciativa_data['linea_url']
    # Síntesis de innovación
    sintesis_innovacion_data = sintesis_Innovación(request)
    tabla_innovacion_url = sintesis_innovacion_data['tabla_url']
    linea_innovacion_url = sintesis_innovacion_data['linea_url']
    # Síntesis de Libertad de pensamiento
    sintesis_libertad_data = sintesis_Libertad_de_pensamiento(request)
    tabla_libertad_url = sintesis_libertad_data['tabla_url']
    linea_libertad_url = sintesis_libertad_data['linea_url']
    # Síntesis de motivación
    sintesis_motivacion_data = sintesis_Motivación(request)
    tabla_motivacion_url = sintesis_motivacion_data['tabla_url']
    linea_motivacion_url = sintesis_motivacion_data['linea_url']
    # Síntesis de pertenencia
    sintesis_pertenencia_data = sintesis_Pertenencia(request)
    tabla_pertenencia_url = sintesis_pertenencia_data['tabla_url']
    linea_pertenencia_url = sintesis_pertenencia_data['linea_url']
    # Síntesis de Satisfacción/Insatisfacción con la organización
    sintesis_satisfaccion_data = sintesis_satisfaccion_insatisfaccion_con_la_organizacion(request)
    tabla_satisfaccion_url = sintesis_satisfaccion_data['tabla_url']
    linea_satisfaccion_url = sintesis_satisfaccion_data['linea_url']
    # Síntesis de "La organización aprovecha las competencias individuales"
    sintesis_organizacion_data = sintesis_La_organizacion_aprovecha_las_competencias_individuales(request)
    tabla_organizacion_url = sintesis_organizacion_data['tabla_url']
    linea_organizacion_url = sintesis_organizacion_data['linea_url']
        # Síntesis de Equidad en la carga laboral
    sintesis_equidad_data = sintesis_Equidad(request)
    tabla_equidad_url = sintesis_equidad_data['tabla_url']
    linea_equidad_url = sintesis_equidad_data['linea_url']
    # Síntesis de Percepción de la normatividad
    sintesis_normatividad_data = sintesis_Percepción_de_la_normatividad(request)
    tabla_normatividad_url = sintesis_normatividad_data['tabla_url']
    linea_normatividad_url = sintesis_normatividad_data['linea_url']
    # Llamada a la síntesis general que consolida los resultados de varias dimensiones.
    sintesis_general_data = sintesis_general(request)
    tabla_sintesis_url = sintesis_general_data['tabla_url']
    linea_sintesis_url = sintesis_general_data['linea_url']










    # Renderizar la plantilla HTML con las URLs de las imágenes
    context = {
        'age_chart_url': age_chart_url,
        'gender_chart_url': gender_chart_url,
        'area_empresa_chart_url': area_empresa_chart_url,
        'antiguedad_empresa_chart_url': antiguedad_empresa_chart_url,
        'data_table_url': table_url,
        'pregunta_1_chart_url': pregunta_1_chart_url,
        'pregunta_2_chart_url': pregunta_2_chart_url,
        'pregunta_3_chart_url': pregunta_3_chart_url,
        'pregunta_4_chart_url': pregunta_4_chart_url,
        'pregunta_5_chart_url': pregunta_5_chart_url,
        'pregunta_6_chart_url': pregunta_6_chart_url,
        'pregunta_7_chart_url': pregunta_7_chart_url,
        'pregunta_8_chart_url': pregunta_8_chart_url,
        'pregunta_9_chart_url': pregunta_9_chart_url,
        'pregunta_10_chart_url': pregunta_10_chart_url,
        'pregunta_11_chart_url': pregunta_11_chart_url,
        'pregunta_12_chart_url': pregunta_12_chart_url,
        'pregunta_13_chart_url': pregunta_13_chart_url,
        'pregunta_14_chart_url': pregunta_14_chart_url,
        'pregunta_15_chart_url': pregunta_15_chart_url,
        'pregunta_16_chart_url': pregunta_16_chart_url,
        'pregunta_17_chart_url': pregunta_17_chart_url,
        'pregunta_18_chart_url': pregunta_18_chart_url,
        'pregunta_19_chart_url': pregunta_19_chart_url,
        'pregunta_20_chart_url': pregunta_20_chart_url,
        'pregunta_21_chart_url': pregunta_21_chart_url,
        'pregunta_22_chart_url': pregunta_22_chart_url,
        'pregunta_23_chart_url': pregunta_23_chart_url,
        'pregunta_24_chart_url': pregunta_24_chart_url,
        'pregunta_25_chart_url': pregunta_25_chart_url,
        'pregunta_26_chart_url': pregunta_26_chart_url,
        'pregunta_27_chart_url': pregunta_27_chart_url,
        'pregunta_28_chart_url': pregunta_28_chart_url,
        'pregunta_29_chart_url': pregunta_29_chart_url,
        'pregunta_30_chart_url': pregunta_30_chart_url,
        'pregunta_31_chart_url': pregunta_31_chart_url,
        'pregunta_32_chart_url': pregunta_32_chart_url,
        'pregunta_33_chart_url': pregunta_33_chart_url,
        'pregunta_34_chart_url': pregunta_34_chart_url,
        'pregunta_35_chart_url': pregunta_35_chart_url,
        'pregunta_36_chart_url': pregunta_36_chart_url,
        'pregunta_37_chart_url': pregunta_37_chart_url,
        'pregunta_38_chart_url': pregunta_38_chart_url,
        'pregunta_39_chart_url': pregunta_39_chart_url,
        'pregunta_40_chart_url': pregunta_40_chart_url,
        'pregunta_41_chart_url': pregunta_41_chart_url,
        'pregunta_42_chart_bar3_url' : pregunta_42_chart_bar3_url,
        'pregunta_43_opcion_1_chart_url': pregunta_43_opcion_1_chart_url,
        'pregunta_43_opcion_2_chart_url': pregunta_43_opcion_2_chart_url,
        'pregunta_43_opcion_3_chart_url': pregunta_43_opcion_3_chart_url,
        'pregunta_43_chart_url': pregunta_43_chart_url,
        'pregunta_44_opcion_1_chart_url': pregunta_44_opcion_1_chart_url,
        'pregunta_44_opcion_2_chart_url': pregunta_44_opcion_2_chart_url,
        'pregunta_44_opcion_3_chart_url': pregunta_44_opcion_3_chart_url,
        'pregunta_44_opcion_4_chart_url': pregunta_44_opcion_4_chart_url,
        'pregunta_44_chart_url': pregunta_44_chart_url,
        'pregunta_45_chart_url': pregunta_45_chart_url,
        'pregunta_46_chart_url': pregunta_46_chart_url,
        'pregunta_47_chart_url': pregunta_47_chart_url,
        'pregunta_48_chart_url': pregunta_48_chart_url,
        'pregunta_49_chart_url': pregunta_49_chart_url,
        'pregunta_50_chart_url': pregunta_50_chart_url,
        'pregunta_51_chart_url': pregunta_51_chart_url,
        'preguntas_52_a_61_table_url' : preguntas_52_a_61_table_url,
        # Síntesis de administración
        'sintesis_tabla_url': tabla_admin_url,
        'sintesis_linea_url': linea_admin_url,
        # Síntesis de calidad
        'sintesis_tabla_calidad_url': tabla_calidad_url,
        'sintesis_linea_calidad_url': linea_calidad_url,
        # Síntesis de cambio
        'sintesis_tabla_cambio_url': tabla_cambio_url,
        'sintesis_linea_cambio_url': linea_cambio_url,
        # Síntesis de comunicación
        'sintesis_tabla_comunicacion_url': tabla_comunicacion_url,
        'sintesis_linea_comunicacion_url': linea_comunicacion_url,
        # Síntesis de conflicto y negociación
        'sintesis_tabla_conflicto_negociacion_url': tabla_conflicto_negociacion_url,
        'sintesis_linea_conflicto_negociacion_url': linea_conflicto_negociacion_url,
        # Síntesis de delegación y confianza
        'sintesis_tabla_delegacion_confianza_url': tabla_delegacion_confianza_url,
        'sintesis_linea_delegacion_confianza_url': linea_delegacion_confianza_url,
        # Síntesis de Eficacia/Coherencia
        'sintesis_tabla_eficacia_coherencia_url': tabla_eficacia_coherencia_url,
        'sintesis_linea_eficacia_coherencia_url': linea_eficacia_coherencia_url,
        # Síntesis de Eficiencia: Ahorro Vs Desperdicio interno
        'sintesis_tabla_eficiencia_url': tabla_eficiencia_url,
        'sintesis_linea_eficiencia_url': linea_eficiencia_url,
        # Síntesis de Liderazgo
        'sintesis_tabla_liderazgo_url': tabla_liderazgo_url,
        'sintesis_linea_liderazgo_url': linea_liderazgo_url,
        # Síntesis de Orientación
        'sintesis_tabla_orientacion_url': sintesis_orientacion_tabla_url,
        'sintesis_linea_orientacion_url': sintesis_orientacion_linea_url,
        # Síntesis de Poder/Autoridad
        'sintesis_tabla_poder_autoridad_url': tabla_poder_autoridad_url,
        'sintesis_linea_poder_autoridad_url': linea_poder_autoridad_url,
        # Síntesis de Trabajo en Equipo
        'sintesis_tabla_trabajo_equipo_url': tabla_trabajo_equipo_url,
        'sintesis_linea_trabajo_equipo_url': linea_trabajo_equipo_url,
        # Síntesis de Calidez/Frialdad
        'sintesis_tabla_calidez_frialdad_url': tabla_calidez_frialdad_url,
        'sintesis_linea_calidez_frialdad_url': linea_calidez_frialdad_url,
        # Síntesis de Autonomía/Control
        'sintesis_tabla_autonomia_control_url': tabla_autonomia_control_url,
        'sintesis_linea_autonomia_control_url': linea_autonomia_control_url,
        # Síntesis de Desarrollo de Carrera
        'sintesis_tabla_desarrollo_url': tabla_desarrollo_url,
        'sintesis_linea_desarrollo_url': linea_desarrollo_url,
        # Síntesis de iniciativa
        'sintesis_tabla_iniciativa_url': tabla_iniciativa_url,
        'sintesis_linea_iniciativa_url': linea_iniciativa_url,
        # Síntesis de innovación
        'sintesis_tabla_innovacion_url': tabla_innovacion_url,
        'sintesis_linea_innovacion_url': linea_innovacion_url,
        # Síntesis de Libertad de pensamiento
        'sintesis_tabla_libertad_url': tabla_libertad_url,
        'sintesis_linea_libertad_url': linea_libertad_url,
        # Síntesis de motivación
        'sintesis_tabla_motivacion_url': tabla_motivacion_url,
        'sintesis_linea_motivacion_url': linea_motivacion_url,
        # Síntesis de pertenencia
        'sintesis_tabla_pertenencia_url': tabla_pertenencia_url,
        'sintesis_linea_pertenencia_url': linea_pertenencia_url,
        # Síntesis de Satisfacción/Insatisfacción con la organización
        'sintesis_tabla_satisfaccion_url': tabla_satisfaccion_url,
        'sintesis_linea_satisfaccion_url': linea_satisfaccion_url,
        # Síntesis de "La organización aprovecha las competencias individuales"
        'sintesis_tabla_organizacion_url': tabla_organizacion_url,
        'sintesis_linea_organizacion_url': linea_organizacion_url,
        # Síntesis de Equidad en la carga laboral
        'sintesis_tabla_equidad_url': tabla_equidad_url,
        'sintesis_linea_equidad_url': linea_equidad_url,
        # Síntesis de Percepción de la normatividad
        'sintesis_tabla_normatividad_url': tabla_normatividad_url,
        'sintesis_linea_normatividad_url': linea_normatividad_url,
        # Resultados de la síntesis general.
        'sintesis_tabla_general_url': tabla_sintesis_url,
        'sintesis_linea_general_url': linea_sintesis_url,
    }
    return render(request, 'results.html', context)
